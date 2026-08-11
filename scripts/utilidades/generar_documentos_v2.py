"""
Genera documentación completa Rossmix v2:
  - docs/Rossmix_CasosDeUso_v2.docx   (Word  — 10 CU con tablas de color)
  - docs/Rossmix_CasosDeUso_v2.xlsx   (Excel — 10 HU + 10 CU con bloques por sección)
"""
import os
from datetime import date

# ════════════════════════════════════════════════════════════════════════════
# DATOS — 10 CASOS DE USO
# ════════════════════════════════════════════════════════════════════════════
CASOS_DE_USO = [
    {
        "id": "CU-01",
        "nombre": "Consultar Servicios y Disponibilidad",
        "actores": "Cliente",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "El cliente explora los servicios ofrecidos, sus precios, "
            "el personal que los realiza y los horarios libres para agendar."
        ),
        "precondiciones": (
            "El administrador debe haber registrado previamente los servicios, "
            "el personal y sus horarios de trabajo."
        ),
        "postcondiciones": (
            "El cliente visualiza la disponibilidad real del salón y puede "
            "continuar al agendamiento."
        ),
        "flujo_principal": [
            "El cliente ingresa a la plataforma web de Rossmix.",
            "Selecciona un servicio del catálogo para ver su precio total y duración estimada.",
            "El sistema muestra la lista de profesionales capacitados para realizar dicho servicio.",
            "El cliente selecciona un profesional (o la opción 'Cualquiera disponible').",
            "El sistema despliega un calendario interactivo con las fechas y franjas horarias disponibles.",
        ],
        "flujos_alt": [
            "Sin disponibilidad: Si el profesional no tiene franjas libres en la fecha "
            "seleccionada, el sistema sugiere las fechas más cercanas con disponibilidad.",
        ],
    },
    {
        "id": "CU-02",
        "nombre": "Agendar Cita y Pagar Abono",
        "actores": "Cliente",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "El cliente reserva una fecha y hora específica tras realizar "
            "el pago obligatorio del abono de $5.000 COP."
        ),
        "precondiciones": (
            "El cliente ha seleccionado un servicio, profesional y franja "
            "horaria disponible (Caso de Uso 1)."
        ),
        "postcondiciones": (
            "La cita queda en estado 'Confirmada', el horario bloqueado y "
            "el cliente recibe comprobante por correo/WhatsApp."
        ),
        "flujo_principal": [
            "El cliente hace clic en 'Reservar'.",
            "El sistema solicita sus datos de contacto (Nombre, Teléfono/WhatsApp y Correo electrónico).",
            "El cliente confirma los datos y el sistema lo redirige a la pasarela de pagos.",
            "El cliente efectúa el pago exitoso del abono ($5.000 COP).",
            "La pasarela confirma el pago al sistema.",
            "El sistema cambia el estado de la cita a 'Confirmada', bloquea el horario "
            "en la agenda y envía un comprobante por correo/WhatsApp.",
        ],
        "flujos_alt": [
            "Pago fallido o cancelado: Si la transacción es rechazada, la cita se libera "
            "inmediatamente y se le notifica al cliente para que reintente el pago.",
        ],
    },
    {
        "id": "CU-03",
        "nombre": "Cancelar Cita con Devolución de Dinero",
        "actores": "Cliente",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "El cliente solicita la cancelación de una cita previamente "
            "agendada dentro del marco de las políticas del salón."
        ),
        "precondiciones": "La cita debe estar en estado 'Confirmada'.",
        "postcondiciones": (
            "La cita pasa a estado 'Cancelada'. Si aplica, el abono de "
            "$5.000 COP es reembolsado."
        ),
        "flujo_principal": [
            "El cliente accede a su cita mediante el enlace enviado a su correo/WhatsApp o desde su panel.",
            "El cliente presiona la opción 'Cancelar Cita'.",
            "El sistema evalúa el tiempo restante: Hora Cita - Hora Actual >= 2 horas.",
            "La condición se cumple; el sistema procesa la cancelación de la cita en la agenda.",
            "El sistema gestiona el reembolso automático de los $5.000 COP a través de la pasarela "
            "(o genera una orden de devolución para administración) y notifica al cliente.",
        ],
        "flujos_alt": [
            "Cancelación fuera de tiempo: Si faltan menos de 2 horas para la cita "
            "(Hora Cita - Hora Actual < 2 horas), el sistema bloquea el reembolso, notifica al cliente "
            "que no aplica la devolución del abono de $5.000 COP según los términos del servicio y libera el cupo.",
        ],
    },
    {
        "id": "CU-04",
        "nombre": "Gestionar Horarios y Personal",
        "actores": "Administrador / Empleado",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "Permite configurar qué servicios presta cada profesional, "
            "sus días laborales, turnos y descansos."
        ),
        "precondiciones": "El usuario ha iniciado sesión como Administrador en el panel interno.",
        "postcondiciones": (
            "Los cambios de horario y servicios quedan guardados y la "
            "disponibilidad pública se actualiza automáticamente."
        ),
        "flujo_principal": [
            "El administrador ingresa a la sección 'Gestión de Personal'.",
            "Selecciona un profesional o crea uno nuevo.",
            "Asigna los servicios que esa persona está capacitada para realizar.",
            "Define el horario semanal (ej. Lunes a Sábado de 8:00 AM a 6:00 PM) "
            "y bloquea franjas de almuerzo o días libres.",
            "El sistema guarda los cambios y actualiza automáticamente la matriz "
            "de disponibilidad pública.",
        ],
        "flujos_alt": [
            "Solapamiento con citas existentes: Si el administrador intenta bloquear un horario "
            "o modificar la jornada de un profesional que ya tiene citas confirmadas en ese rango, "
            "el sistema despliega una alerta listando los clientes afectados y solicita confirmación "
            "para cancelar o reprogramar manualmente dichas citas antes de aplicar el cambio.",
            "Sin servicios asignados: Si el administrador guarda la configuración de un profesional "
            "sin asignarle al menos un servicio, el sistema muestra un mensaje de advertencia y no "
            "publica al empleado en la vista del cliente.",
        ],
    },
    {
        "id": "CU-05",
        "nombre": "Visualizar y Gestionar Agenda del Día",
        "actores": "Administrador / Recepcionista",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "Permite ver las citas programadas del día, registrar la llegada "
            "del cliente y saldar el pago restante del servicio."
        ),
        "precondiciones": "El administrador ha iniciado sesión en el backoffice de Rossmix.",
        "postcondiciones": (
            "Las citas del día quedan con estado actualizado y los pagos "
            "del saldo pendiente registrados."
        ),
        "flujo_principal": [
            "El usuario accede a la vista de 'Agenda Diaria'.",
            "El sistema muestra una cuadrícula filtrada por profesionales con las citas del día.",
            "Al llegar un cliente, el usuario marca la cita como 'En atención'.",
            "Una vez terminado el servicio, el sistema calcula el saldo pendiente: "
            "Saldo Pendiente = Precio Total - $5.000 COP.",
            "El usuario registra el pago restante en efectivo o transferencia y marca la cita como 'Completada'.",
        ],
        "flujos_alt": [
            "Inasistencia del cliente (No-Show): Si transcurre un margen de tolerancia predefinido "
            "(ej. 15-20 minutos) tras la hora pactada sin que el cliente se presente, el usuario marca "
            "la cita como 'No asistió'. El sistema libera al profesional, marca el abono de $5.000 COP "
            "como no reembolsable y actualiza el historial del cliente.",
            "Servicios adicionales o cambios en atención: Si durante la cita el cliente solicita un "
            "servicio extra o cambia el acordado, el usuario modifica la cita en pantalla. El sistema "
            "recalcula automáticamente el valor total y ajusta el nuevo saldo pendiente a cobrar.",
        ],
    },
    {
        "id": "CU-06",
        "nombre": "Reprogramar / Modificar Cita",
        "actores": "Cliente / Administrador",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "Permite al cliente cambiar la fecha, la hora o el profesional de una cita "
            "previamente confirmada sin perder su abono de $5.000 COP."
        ),
        "precondiciones": (
            "La cita debe estar en estado 'Confirmada' y debe cumplirse la regla de tiempo "
            "(mínimo 2 horas antes de la hora pactada)."
        ),
        "postcondiciones": (
            "La cita original queda cancelada y se crea una nueva con el abono "
            "transferido. El cliente recibe confirmación."
        ),
        "flujo_principal": [
            "El cliente ingresa a la opción 'Gestionar Cita' desde el enlace de confirmación "
            "recibido en WhatsApp/correo.",
            "El cliente selecciona la opción 'Reprogramar Cita'.",
            "El sistema evalúa si la hora actual cumple con la política: "
            "Hora Cita - Hora Actual >= 2 horas.",
            "El sistema despliega el calendario con la nueva disponibilidad del profesional "
            "o permite elegir uno nuevo.",
            "El cliente selecciona la nueva fecha y hora disponible.",
            "El sistema actualiza los detalles de la cita, mantiene el abono de $5.000 COP "
            "asociado al nuevo registro y envía una actualización de la reserva por correo/WhatsApp.",
        ],
        "flujos_alt": [
            "Intento fuera de tiempo: Si faltan menos de 2 horas para la cita, el sistema notifica "
            "que no es posible reprogramar de manera autónoma y sugiere comunicarse directamente con el salón.",
            "Reprogramación por parte del salón (Admin): Si un profesional presenta un imprevisto, "
            "el administrador puede mover la cita del cliente desde el panel interno notificándose "
            "automáticamente la reprogramación.",
        ],
    },
    {
        "id": "CU-07",
        "nombre": "Gestión de Pagos y Reembolsos",
        "actores": "Administrador / Sistema",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "Registra, confirma y gestiona todos los movimientos de pago del salón: "
            "abonos, saldos pendientes, pagos adicionales y reembolsos."
        ),
        "precondiciones": "Existe una cita registrada con abono o pago pendiente.",
        "postcondiciones": (
            "El pago queda registrado en la tabla pagos. El saldo pendiente "
            "de la cita se actualiza. Si aplica, el reembolso se procesa."
        ),
        "flujo_principal": [
            "El administrador accede al módulo 'Pagos' desde el panel.",
            "Selecciona la cita a liquidar o el pago a registrar.",
            "Ingresa el monto, método de pago (efectivo, tarjeta, transferencia, Nequi, Daviplata) "
            "y referencia opcional.",
            "El sistema valida que monto > 0 y registra el pago en la tabla pagos.",
            "El sistema actualiza el saldo_pendiente de la cita y cambia el estado si corresponde.",
            "El cliente recibe notificación del pago registrado.",
        ],
        "flujos_alt": [
            "Reembolso por cancelación: Si la cita fue cancelada con ≥2 horas de anticipación, "
            "el sistema genera automáticamente el registro de reembolso y marca cita.reembolsado=True.",
            "Pago duplicado: Si ya existe un pago registrado para la cita (UNIQUE en id_cita), "
            "el sistema notifica el error y bloquea el registro.",
        ],
    },
    {
        "id": "CU-08",
        "nombre": "Gestión de Notificaciones",
        "actores": "Sistema / Cliente / Administrador",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "El sistema genera y envía notificaciones internas y por correo/WhatsApp "
            "ante eventos relevantes: confirmación, cancelación, reprogramación, recordatorios."
        ),
        "precondiciones": "Ocurre un evento relevante en el sistema (nueva cita, cambio de estado, pago).",
        "postcondiciones": (
            "La notificación queda registrada en la tabla notificaciones "
            "y el usuario es alertado en la plataforma y/o por email."
        ),
        "flujo_principal": [
            "Se produce un evento en el sistema (cita confirmada, cancelada, reprogramada, pago registrado).",
            "El sistema crea un registro en la tabla notificaciones con título, mensaje y target URL.",
            "El usuario ve el contador de notificaciones no leídas en la barra de navegación.",
            "El usuario hace clic y accede al centro de notificaciones.",
            "Puede marcar una o todas como leídas.",
        ],
        "flujos_alt": [
            "Fallo en envío de correo: Si Flask-Mail no puede conectar al servidor SMTP, "
            "la notificación interna se guarda igualmente; el error se registra en logs pero "
            "no interrumpe el flujo principal.",
            "Notificación masiva: El administrador puede enviar avisos a todos los clientes "
            "con citas en un rango de fechas (ej. cierre por festivo).",
        ],
    },
    {
        "id": "CU-09",
        "nombre": "Auditoría y Seguridad de Usuarios",
        "actores": "Administrador / Sistema",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "Registra automáticamente todas las acciones críticas sobre cuentas de usuario: "
            "login, logout, edición, desactivación y eliminación, con IP y actor."
        ),
        "precondiciones": "El sistema está en operación y un usuario realiza una acción relevante.",
        "postcondiciones": (
            "La acción queda registrada en auditoria_usuarios con timestamp, "
            "IP, actor y snapshot de los datos."
        ),
        "flujo_principal": [
            "Un usuario realiza una acción relevante (login, editar cuenta, eliminar cliente).",
            "El sistema captura: acción, id_usuario afectado, id_actor, IP, timestamp.",
            "Se guarda un snapshot de nombre, email, teléfono y rol al momento del evento.",
            "El registro se persiste en auditoria_usuarios.",
            "El administrador puede consultar el historial de auditoría filtrado por usuario, "
            "acción o rango de fechas.",
        ],
        "flujos_alt": [
            "Cuenta desactivada con sesión activa: Al detectar que usuario.activo=False, "
            "el decorador admin_required/especialista_required limpia la sesión y redirige al login.",
            "Múltiples intentos fallidos: Se registra cada intento fallido de login "
            "con IP para detección de ataques de fuerza bruta.",
        ],
    },
    {
        "id": "CU-10",
        "nombre": "Configuración del Sistema",
        "actores": "Administrador",
        "modulo": "Gestión de Citas Rossmix",
        "descripcion": (
            "Permite al administrador ajustar los parámetros operativos del sistema "
            "(abono mínimo, anticipación mínima, nombre del salón, moneda) sin modificar código."
        ),
        "precondiciones": "Usuario con rol admin con sesión activa.",
        "postcondiciones": (
            "El parámetro queda actualizado en la tabla configuraciones con registro "
            "del usuario que lo modificó y timestamp."
        ),
        "flujo_principal": [
            "El administrador accede al módulo 'Configuración del Sistema'.",
            "Ve la lista de parámetros actuales con su valor y descripción.",
            "Selecciona el parámetro a modificar (ej. abono_minimo).",
            "Ingresa el nuevo valor y confirma.",
            "El sistema actualiza configuraciones.valor, registra modificado_por=admin.id "
            "y fecha_actualizacion=now().",
            "Los cambios se aplican inmediatamente en el flujo de agendamiento.",
        ],
        "flujos_alt": [
            "Valor inválido: Si el administrador ingresa un valor fuera del rango permitido "
            "(ej. abono_minimo < 0), el sistema muestra error de validación sin guardar.",
            "Parámetro crítico: Para parámetros que afectan citas activas (ej. anticipacion_minima), "
            "el sistema solicita confirmación explicando el impacto.",
        ],
    },
]

# ════════════════════════════════════════════════════════════════════════════
# DATOS — 10 HISTORIAS DE USUARIO
# ════════════════════════════════════════════════════════════════════════════
HISTORIAS = [
    {
        "id": "HU-01", "cu": "CU-01", "prioridad": "Alta", "puntos": "5 Pts",
        "titulo": "Consultar Servicios y Disponibilidad",
        "actor_titulo": "CLIENTE DE ROSSMIX",
        "subtitulo": "Especificación ágil de requerimiento para Consultar Servicios y Disponibilidad",
        "rol": "Cliente de Rossmix", "estado": "Por Hacer",
        "quiero": "Explorar el catálogo de servicios con sus precios y ver la disponibilidad de los profesionales",
        "para": "Elegir el servicio y el horario que mejor se adapten a mi tiempo antes de agendar.",
        "criterios": [
            "Debe mostrar duración estimada y precio total de cada servicio.",
            "Debe permitir seleccionar un profesional específico o la opción 'Cualquiera disponible'.",
            "Debe desplegar un calendario en tiempo real con franjas horarias libres.",
            "Si no hay disponibilidad en la fecha seleccionada, debe sugerir las fechas más cercanas.",
        ],
    },
    {
        "id": "HU-02", "cu": "CU-02", "prioridad": "Alta", "puntos": "8 Pts",
        "titulo": "Agendar Cita y Pagar Abono",
        "actor_titulo": "CLIENTE",
        "subtitulo": "Especificación ágil de requerimiento para Agendar Cita y Pagar Abono",
        "rol": "Cliente", "estado": "Por Hacer",
        "quiero": "Reservar mi cita realizando el pago obligatorio del abono de $5.000 COP a través de la plataforma",
        "para": "Asegurar mi cupo en el salón de forma rápida, transparente y confiable.",
        "criterios": [
            "Debe solicitar datos básicos: Nombre, Teléfono/WhatsApp y Correo electrónico.",
            "Debe redirigir a una pasarela de pagos segura para cobrar los $5.000 COP.",
            "Al confirmarse el pago, la cita cambia automáticamente a estado 'Confirmada' y bloquea el horario.",
            "Se debe enviar una notificación/comprobante de reserva por correo electrónico y/o WhatsApp.",
            "Si el pago falla o se cancela, el horario debe liberarse inmediatamente.",
        ],
    },
    {
        "id": "HU-03", "cu": "CU-03", "prioridad": "Media", "puntos": "5 Pts",
        "titulo": "Cancelar Cita con Devolución de Dinero",
        "actor_titulo": "CLIENTE CON CITA CONFIRMADA",
        "subtitulo": "Especificación ágil de requerimiento para Cancelar Cita con Devolución de Dinero",
        "rol": "Cliente con cita confirmada", "estado": "Por Hacer",
        "quiero": "Poder cancelar mi reserva desde la plataforma",
        "para": "Liberar mi turno y solicitar la devolución de mi abono de $5.000 COP si aviso con suficiente anticipación.",
        "criterios": [
            "El cliente puede solicitar la cancelación desde su enlace/panel de gestión.",
            "Si la cancelación ocurre con >= 2 horas de anticipación, el sistema procesa la devolución de los $5.000 COP y libera la agenda.",
            "Si la cancelación ocurre a < 2 horas, el sistema notifica que no aplica devolución según términos y condiciones, pero libera la franja horaria.",
        ],
    },
    {
        "id": "HU-04", "cu": "CU-04", "prioridad": "Alta", "puntos": "5 Pts",
        "titulo": "Gestionar Horarios y Personal",
        "actor_titulo": "ADMINISTRADOR DE ROSSMIX",
        "subtitulo": "Especificación ágil de requerimiento para Gestionar Horarios y Personal",
        "rol": "Administrador de Rossmix", "estado": "Por Hacer",
        "quiero": "Configurar la agenda laboral, turnos y servicios asignados a cada profesional",
        "para": "Mantener actualizada la disponibilidad del salón y evitar cruces de horarios.",
        "criterios": [
            "Permite asignar qué servicios puede realizar cada empleado.",
            "Permite establecer horarios semanales, descansos y días festivos o libres.",
            "Si se intenta bloquear un horario donde ya existen citas agendadas, el sistema debe alertar para reprogramar o cancelar a los clientes afectados.",
            "No se debe publicar en la vista del cliente a ningún profesional que no tenga al menos un servicio asociado.",
        ],
    },
    {
        "id": "HU-05", "cu": "CU-05", "prioridad": "Alta", "puntos": "8 Pts",
        "titulo": "Visualizar y Gestionar Agenda del Día",
        "actor_titulo": "RECEPCIONISTA / ADMINISTRADOR",
        "subtitulo": "Especificación ágil de requerimiento para Visualizar y Gestionar Agenda del Día",
        "rol": "Recepcionista / Administrador", "estado": "Por Hacer",
        "quiero": "Visualizar la cuadrícula diaria de citas y registrar el pago del saldo pendiente",
        "para": "Controlar la atención en tiempo real y saldar las cuentas de cada servicio completado.",
        "criterios": [
            "Interfaz con vista en cuadrícula de las citas del día organizadas por profesional.",
            "Permite cambiar estados de cita ('En atención', 'Completada', 'No asistió').",
            "Calcula automáticamente el saldo restante: Saldo Pendiente = Precio Total - $5.000 COP.",
            "Permite agregar servicios adicionales durante la atención y recalcula el monto final a cobrar.",
            "Al marcar 'No asistió', libera la agenda del empleado y registra la inasistencia.",
        ],
    },
    {
        "id": "HU-06", "cu": "CU-06", "prioridad": "Media", "puntos": "5 Pts",
        "titulo": "Reprogramar / Modificar Cita",
        "actor_titulo": "CLIENTE O ADMINISTRADOR",
        "subtitulo": "Especificación ágil de requerimiento para Reprogramar / Modificar Cita",
        "rol": "Cliente o Administrador", "estado": "Por Hacer",
        "quiero": "Cambiar la fecha u hora de una cita agendada",
        "para": "Ajustar la atención a un nuevo horario sin perder el abono de $5.000 COP previamente realizado.",
        "criterios": [
            "Si el cliente reprograma con >= 2 horas de anticipación, el sistema le permite seleccionar una nueva fecha/hora manteniendo el abono.",
            "Si faltan menos de 2 horas, el sistema bloquea la edición autónoma y le indica contactar directamente al salón.",
            "El administrador puede reprogramar citas de cualquier cliente desde el panel (ej. ante imprevistos del personal) notificando al cliente de forma automática.",
        ],
    },
    {
        "id": "HU-07", "cu": "CU-07", "prioridad": "Alta", "puntos": "5 Pts",
        "titulo": "Gestión de Pagos y Reembolsos",
        "actor_titulo": "ADMINISTRADOR DE ROSSMIX",
        "subtitulo": "Especificación ágil de requerimiento para Gestión de Pagos y Reembolsos",
        "rol": "Administrador de Rossmix", "estado": "Por Hacer",
        "quiero": "Registrar y gestionar todos los movimientos de pago del salón desde el panel",
        "para": "Llevar control financiero preciso de abonos, saldos cobrados y reembolsos procesados.",
        "criterios": [
            "Permite registrar pagos con método (efectivo, tarjeta, transferencia, Nequi, Daviplata) y referencia.",
            "El sistema actualiza automáticamente el saldo_pendiente de la cita tras registrar el pago.",
            "Los reembolsos de cancelaciones con >= 2h se procesan y registran automáticamente.",
            "No permite registrar dos pagos para la misma cita (UNIQUE constraint).",
            "El admin puede exportar el historial de pagos a Excel por período.",
        ],
    },
    {
        "id": "HU-08", "cu": "CU-08", "prioridad": "Media", "puntos": "3 Pts",
        "titulo": "Gestión de Notificaciones",
        "actor_titulo": "CLIENTE / ADMINISTRADOR",
        "subtitulo": "Especificación ágil de requerimiento para Gestión de Notificaciones",
        "rol": "Cliente / Administrador", "estado": "Por Hacer",
        "quiero": "Recibir alertas en tiempo real de los eventos relevantes de mis citas",
        "para": "Estar siempre informado sobre confirmaciones, cambios y recordatorios sin tener que revisar manualmente.",
        "criterios": [
            "El sistema genera notificación interna y por email ante: confirmación, cancelación, reprogramación y pago.",
            "Las notificaciones no leídas muestran contador en la barra de navegación.",
            "El usuario puede marcar notificaciones individuales o todas como leídas.",
            "Si el envío de correo falla, la notificación interna se guarda igualmente.",
        ],
    },
    {
        "id": "HU-09", "cu": "CU-09", "prioridad": "Alta", "puntos": "5 Pts",
        "titulo": "Auditoría y Seguridad de Usuarios",
        "actor_titulo": "ADMINISTRADOR DE ROSSMIX",
        "subtitulo": "Especificación ágil de requerimiento para Auditoría y Seguridad",
        "rol": "Administrador de Rossmix", "estado": "Por Hacer",
        "quiero": "Consultar el historial de acciones críticas realizadas sobre las cuentas de usuario",
        "para": "Garantizar la trazabilidad y seguridad del sistema ante cambios no autorizados.",
        "criterios": [
            "Registra automáticamente: login, logout, edición, desactivación y eliminación de cuentas.",
            "Cada registro incluye: acción, usuario afectado, admin actor, IP y timestamp.",
            "Si una cuenta es desactivada, su sesión activa se invalida inmediatamente.",
            "El admin puede filtrar el historial por usuario, acción o rango de fechas.",
        ],
    },
    {
        "id": "HU-10", "cu": "CU-10", "prioridad": "Media", "puntos": "3 Pts",
        "titulo": "Configuración del Sistema",
        "actor_titulo": "ADMINISTRADOR DE ROSSMIX",
        "subtitulo": "Especificación ágil de requerimiento para Configuración del Sistema",
        "rol": "Administrador de Rossmix", "estado": "Por Hacer",
        "quiero": "Modificar los parámetros operativos del sistema desde el panel sin tocar el código",
        "para": "Adaptar el comportamiento del salón (abono, anticipación, nombre) sin depender de un desarrollador.",
        "criterios": [
            "Permite editar parámetros clave: abono_minimo, anticipacion_minima, max_dias_agenda, nombre_salon, moneda.",
            "Cada cambio registra quién lo modificó y cuándo (modificado_por + fecha_actualizacion).",
            "Los cambios se aplican en tiempo real sin reiniciar el servidor.",
            "Valores inválidos son rechazados con mensaje de error descriptivo.",
        ],
    },
]

# ════════════════════════════════════════════════════════════════════════════
# GENERAR WORD
# ════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_ROJO   = 'C41E3A'
C_MORADO = '5B21B6'
C_VERDE  = '065F46'
C_AZUL   = '1E3A5F'
C_GRIS   = 'F1F5F9'

def rgb(hex6):
    return RGBColor(int(hex6[0:2],16), int(hex6[2:4],16), int(hex6[4:6],16))

def cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

def add_cu_table(doc, cu):
    """Crea la tabla de un Caso de Uso con el formato de las imágenes."""
    # Título del CU
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{cu['id']}: {cu['nombre']}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = rgb(C_ROJO)

    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'

    def fila(label, valor, bg_label=C_AZUL, bold_label=True):
        row = t.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = label
        c1.text = valor
        cell_bg(c0, bg_label)
        for run in c0.paragraphs[0].runs:
            run.bold = bold_label
            run.font.color.rgb = rgb('FFFFFF')
            run.font.size = Pt(9)
        for run in c1.paragraphs[0].runs:
            run.font.size = Pt(9)
        c0.width = Cm(4.5)
        c1.width = Cm(12)

    fila('ID-CU',        cu['id'])
    fila('Nombre CU',    cu['nombre'])
    fila('Descripción',  cu['descripcion'])
    fila('Actores',      cu['actores'])
    fila('Precondiciones', cu['precondiciones'])

    # Flujo principal
    row_fp = t.add_row()
    cell_bg(row_fp.cells[0], C_AZUL)
    r = row_fp.cells[0].paragraphs[0].add_run('Flujo principal')
    r.bold = True; r.font.color.rgb = rgb('FFFFFF'); r.font.size = Pt(9)
    fp_text = '\n'.join(f"{i+1}.{s}" for i, s in enumerate(cu['flujo_principal']))
    row_fp.cells[1].text = fp_text
    for run in row_fp.cells[1].paragraphs[0].runs:
        run.font.size = Pt(9)

    # Flujos alternativos
    row_fa = t.add_row()
    cell_bg(row_fa.cells[0], C_AZUL)
    r2 = row_fa.cells[0].paragraphs[0].add_run('Flujos Alternativos /\nExcepciones')
    r2.bold = True; r2.font.color.rgb = rgb('FFFFFF'); r2.font.size = Pt(9)
    fa_text = '\n'.join(cu['flujos_alt'])
    row_fa.cells[1].text = fa_text
    for run in row_fa.cells[1].paragraphs[0].runs:
        run.font.size = Pt(9)

    doc.add_paragraph()

def generar_word():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── Portada ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CASOS DE USO')
    run.bold = True; run.font.size = Pt(20)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('ROSSMIX — Salón de Belleza')
    r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = rgb(C_ROJO)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f'Sistema de Agendamiento de Citas  ·  v2.0  ·  {date.today().strftime("%d/%m/%Y")}')
    r3.font.size = Pt(11); r3.font.color.rgb = rgb('64748B')

    doc.add_paragraph()
    doc.add_paragraph()

    # Tabla resumen de actores
    p4 = doc.add_paragraph()
    r4 = p4.add_run('Actores del Sistema')
    r4.bold = True; r4.font.size = Pt(12); r4.font.color.rgb = rgb(C_ROJO)

    ta = doc.add_table(rows=1, cols=2)
    ta.style = 'Table Grid'
    hdr = ta.rows[0].cells
    hdr[0].text = 'Actor'
    hdr[1].text = 'Descripción'
    for c in hdr:
        cell_bg(c, C_AZUL)
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = rgb('FFFFFF'); run.font.size = Pt(9)
    for actor, desc in [
        ('Cliente',          'Agenda, cancela y reprograma sus propias citas.'),
        ('Administrador',    'Gestión completa: empleados, servicios, pagos, reportes y configuración.'),
        ('Recepcionista',    'Gestiona la agenda diaria y registra pagos de saldo pendiente.'),
        ('Especialista',     'Ve sus citas asignadas, acepta disponibles y actualiza estados.'),
        ('Sistema',          'Procesa notificaciones, auditoría, reembolsos y tokens automáticamente.'),
    ]:
        row = ta.add_row().cells
        row[0].text = actor; row[1].text = desc
        for c in row:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
    ta.columns[0].width = Cm(4)
    ta.columns[1].width = Cm(12)

    doc.add_page_break()

    # ── Los 10 Casos de Uso ───────────────────────────────────────────────────
    for i, cu in enumerate(CASOS_DE_USO):
        add_cu_table(doc, cu)
        if i % 2 == 1:  # salto de página cada 2 CU
            doc.add_page_break()

    os.makedirs('docs', exist_ok=True)
    out = 'docs/Rossmix_CasosDeUso_v2.docx'
    doc.save(out)
    print(f'[OK] Word: {out}')

# ════════════════════════════════════════════════════════════════════════════
# GENERAR EXCEL — formato exacto de las imágenes
# ════════════════════════════════════════════════════════════════════════════
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

def xl_fill(hex6):   return PatternFill('solid', fgColor=hex6)
def xl_font(bold=False, size=11, color='000000', italic=False):
    return Font(bold=bold, size=size, color=color, italic=italic, name='Calibri')
def xl_align(h='left', v='center', wrap=False):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)
def xl_border(color='D1D5DB'):
    s = Side(style='thin', color=color)
    return Border(left=s, right=s, top=s, bottom=s)
def xl_border_none():
    return Border()

FILL_TITULO  = xl_fill('C41E3A')   # rojo rossmix — fila título HU/CU
FILL_HDR_SEC = xl_fill('E2E8F0')   # gris azulado — encabezados de sección
FILL_LABEL   = xl_fill('F8FAFC')   # gris muy claro — celdas label (COMO/QUIERO/PARA)
FILL_AZUL    = xl_fill('1E3A5F')   # azul oscuro — sección INFO
FILL_WHITE   = xl_fill('FFFFFF')

def write_hu_block(ws, hu, start_row):
    """Escribe un bloque Historia de Usuario como en las imágenes."""
    r = start_row

    # ── Fila 1: Título grande ─────────────────────────────────────────────────
    ws.row_dimensions[r].height = 30
    c = ws.cell(row=r, column=1, value=f'HISTORIA DE USUARIO: {hu["actor_titulo"]}')
    c.font    = xl_font(bold=True, size=16, color='C41E3A')
    c.fill    = FILL_WHITE
    c.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    r += 1

    # ── Fila 2: Subtítulo ─────────────────────────────────────────────────────
    ws.row_dimensions[r].height = 18
    c2 = ws.cell(row=r, column=1, value=hu['subtitulo'])
    c2.font = xl_font(italic=True, size=10, color='64748B')
    c2.fill = FILL_WHITE
    c2.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    r += 1

    # ── Fila vacía ────────────────────────────────────────────────────────────
    ws.row_dimensions[r].height = 8
    r += 1

    # ── Sección INFORMACIÓN GENERAL ──────────────────────────────────────────
    ws.row_dimensions[r].height = 20
    sec = ws.cell(row=r, column=1, value='INFORMACIÓN GENERAL')
    sec.font = xl_font(bold=True, size=10, color='1E293B')
    sec.fill = FILL_HDR_SEC
    sec.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    sec.border = xl_border('CBD5E1')
    r += 1

    # Fila: ID Requerimiento | valor | Prioridad | valor
    ws.row_dimensions[r].height = 18
    for col, txt, bold in [(1,'ID Requerimiento:',True),(2,hu['id'],False),(3,'Prioridad:',True),(4,hu['prioridad'],False)]:
        c = ws.cell(row=r, column=col, value=txt)
        c.font = xl_font(bold=bold, size=10)
        c.fill = FILL_WHITE
        c.alignment = xl_align('left', 'center')
        c.border = xl_border()
    r += 1

    # Fila: Título | valor | Estimación | valor
    ws.row_dimensions[r].height = 18
    for col, txt, bold in [(1,'Título:',True),(2,hu['titulo'],False),(3,'Estimación:',True),(4,hu['puntos'],False)]:
        c = ws.cell(row=r, column=col, value=txt)
        c.font = xl_font(bold=bold, size=10)
        c.fill = FILL_WHITE
        c.alignment = xl_align('left', 'center')
        c.border = xl_border()
    r += 1

    # Fila: Actor/Rol | valor | Estado | valor
    ws.row_dimensions[r].height = 18
    for col, txt, bold in [(1,'Actor / Rol:',True),(2,hu['rol'],False),(3,'Estado Actual:',True),(4,hu['estado'],False)]:
        c = ws.cell(row=r, column=col, value=txt)
        c.font = xl_font(bold=bold, size=10)
        c.fill = FILL_WHITE
        c.alignment = xl_align('left', 'center')
        c.border = xl_border()
    r += 1

    ws.row_dimensions[r].height = 8
    r += 1

    # ── Sección DESCRIPCIÓN ÁGIL ─────────────────────────────────────────────
    ws.row_dimensions[r].height = 20
    sec2 = ws.cell(row=r, column=1, value='DESCRIPCIÓN DE LA HISTORIA (FORMATO ÁGIL)')
    sec2.font = xl_font(bold=True, size=10, color='1E293B')
    sec2.fill = FILL_HDR_SEC
    sec2.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    sec2.border = xl_border('CBD5E1')
    r += 1

    for label, valor in [('COMO:', hu['rol']), ('QUIERO:', hu['quiero']), ('PARA:', hu['para'])]:
        ws.row_dimensions[r].height = 22
        lbl = ws.cell(row=r, column=1, value=label)
        lbl.font = xl_font(bold=True, size=10, color='C41E3A')
        lbl.fill = FILL_WHITE
        lbl.alignment = xl_align('left', 'center')
        lbl.border = xl_border()

        val = ws.cell(row=r, column=2, value=valor)
        val.font = xl_font(size=10)
        val.fill = FILL_WHITE
        val.alignment = xl_align('left', 'center', wrap=True)
        val.border = xl_border()
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
        r += 1

    ws.row_dimensions[r].height = 8
    r += 1

    # ── Sección CRITERIOS ────────────────────────────────────────────────────
    ws.row_dimensions[r].height = 20
    sec3 = ws.cell(row=r, column=1, value='CRITERIOS DE ACEPTACIÓN')
    sec3.font = xl_font(bold=True, size=10, color='1E293B')
    sec3.fill = FILL_HDR_SEC
    sec3.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    sec3.border = xl_border('CBD5E1')
    r += 1

    for crit in hu['criterios']:
        ws.row_dimensions[r].height = 22
        c = ws.cell(row=r, column=1, value=f'• {crit}')
        c.font = xl_font(size=10)
        c.fill = FILL_WHITE
        c.alignment = xl_align('left', 'center', wrap=True)
        c.border = xl_border()
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
        r += 1

    ws.row_dimensions[r].height = 15
    r += 2  # separador entre bloques

    return r


def write_cu_block(ws, cu, start_row):
    """Escribe un bloque Caso de Uso como en las imágenes del Excel."""
    r = start_row

    # ── Fila 1: Título ────────────────────────────────────────────────────────
    ws.row_dimensions[r].height = 30
    c = ws.cell(row=r, column=1, value=f'CASO DE USO: {cu["id"]}')
    c.font = xl_font(bold=True, size=16, color='1E3A5F')
    c.fill = FILL_WHITE
    c.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    r += 1

    ws.row_dimensions[r].height = 18
    c2 = ws.cell(row=r, column=1, value=f'Especificación técnica detallada para {cu["nombre"]}')
    c2.font = xl_font(italic=True, size=10, color='64748B')
    c2.fill = FILL_WHITE
    c2.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    r += 1

    ws.row_dimensions[r].height = 8
    r += 1

    # ── INFORMACIÓN DEL CASO DE USO ──────────────────────────────────────────
    ws.row_dimensions[r].height = 20
    sec = ws.cell(row=r, column=1, value='INFORMACIÓN DEL CASO DE USO')
    sec.font = xl_font(bold=True, size=10, color='1E293B')
    sec.fill = FILL_HDR_SEC
    sec.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    sec.border = xl_border('CBD5E1')
    r += 1

    for (l1, v1, l2, v2) in [
        ('ID Caso de Uso:', cu['id'],     'Actor Principal:', cu['actores']),
        ('Nombre CU:',      cu['nombre'], 'Módulo:',          cu['modulo']),
    ]:
        ws.row_dimensions[r].height = 18
        for col, txt, bold in [(1,l1,True),(2,v1,False),(3,l2,True),(4,v2,False)]:
            c = ws.cell(row=r, column=col, value=txt)
            c.font = xl_font(bold=bold, size=10)
            c.fill = FILL_WHITE
            c.alignment = xl_align('left', 'center', wrap=True)
            c.border = xl_border()
        r += 1

    ws.row_dimensions[r].height = 8
    r += 1

    # ── DESCRIPCIÓN Y PRECONDICIONES ─────────────────────────────────────────
    ws.row_dimensions[r].height = 20
    sec2 = ws.cell(row=r, column=1, value='DESCRIPCIÓN Y PRECONDICIONES')
    sec2.font = xl_font(bold=True, size=10, color='1E293B')
    sec2.fill = FILL_HDR_SEC
    sec2.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    sec2.border = xl_border('CBD5E1')
    r += 1

    for label, valor in [('Descripción:', cu['descripcion']), ('Precondiciones:', cu['precondiciones'])]:
        ws.row_dimensions[r].height = 28
        lbl = ws.cell(row=r, column=1, value=label)
        lbl.font = xl_font(bold=True, size=10)
        lbl.fill = FILL_WHITE
        lbl.alignment = xl_align('left', 'top')
        lbl.border = xl_border()

        val = ws.cell(row=r, column=2, value=valor)
        val.font = xl_font(size=10)
        val.fill = FILL_WHITE
        val.alignment = xl_align('left', 'top', wrap=True)
        val.border = xl_border()
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
        r += 1

    ws.row_dimensions[r].height = 8
    r += 1

    # ── FLUJO PRINCIPAL ──────────────────────────────────────────────────────
    ws.row_dimensions[r].height = 20
    sec3 = ws.cell(row=r, column=1, value='FLUJO PRINCIPAL DE PASOS')
    sec3.font = xl_font(bold=True, size=10, color='1E293B')
    sec3.fill = FILL_HDR_SEC
    sec3.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    sec3.border = xl_border('CBD5E1')
    r += 1

    for i, paso in enumerate(cu['flujo_principal']):
        ws.row_dimensions[r].height = 22
        c = ws.cell(row=r, column=1, value=f'{i+1}. {paso}')
        c.font = xl_font(size=10)
        c.fill = FILL_WHITE
        c.alignment = xl_align('left', 'top', wrap=True)
        c.border = xl_border()
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
        r += 1

    ws.row_dimensions[r].height = 8
    r += 1

    # ── FLUJOS ALTERNATIVOS ──────────────────────────────────────────────────
    ws.row_dimensions[r].height = 20
    sec4 = ws.cell(row=r, column=1, value='FLUJOS ALTERNATIVOS Y EXCEPCIONES')
    sec4.font = xl_font(bold=True, size=10, color='1E293B')
    sec4.fill = FILL_HDR_SEC
    sec4.alignment = xl_align('left', 'center')
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    sec4.border = xl_border('CBD5E1')
    r += 1

    for alt in cu['flujos_alt']:
        ws.row_dimensions[r].height = 28
        c = ws.cell(row=r, column=1, value=f'• {alt}')
        c.font = xl_font(size=10)
        c.fill = FILL_WHITE
        c.alignment = xl_align('left', 'top', wrap=True)
        c.border = xl_border()
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
        r += 1

    ws.row_dimensions[r].height = 15
    r += 2

    return r


def generar_excel():
    wb = openpyxl.Workbook()

    # ── Anchos de columna comunes ─────────────────────────────────────────────
    def set_cols(ws):
        ws.sheet_view.showGridLines = False
        widths = [22, 30, 18, 20, 14, 14]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    # ── Hoja índice (primera hoja) ────────────────────────────────────────────
    ws_idx = wb.active
    ws_idx.title = 'Índice'
    ws_idx.sheet_view.showGridLines = False
    ws_idx.column_dimensions['A'].width = 12
    ws_idx.column_dimensions['B'].width = 40
    ws_idx.column_dimensions['C'].width = 20

    ws_idx.row_dimensions[1].height = 35
    t = ws_idx.cell(row=1, column=1, value='ROSSMIX — Documentación del Sistema')
    t.font = xl_font(bold=True, size=16, color='C41E3A')
    t.alignment = xl_align('left', 'center')
    ws_idx.merge_cells('A1:C1')

    ws_idx.row_dimensions[2].height = 18
    s = ws_idx.cell(row=2, column=1, value=f'Versión 2.0  ·  {date.today().strftime("%d/%m/%Y")}  ·  10 Casos de Uso  ·  10 Historias de Usuario')
    s.font = xl_font(italic=True, size=10, color='64748B')
    s.alignment = xl_align('left', 'center')
    ws_idx.merge_cells('A2:C2')

    ws_idx.row_dimensions[4].height = 20
    for ci, hdr in enumerate(['Hoja', 'Contenido', 'Tipo'], 1):
        c = ws_idx.cell(row=4, column=ci, value=hdr)
        c.font = xl_font(bold=True, size=10, color='FFFFFF')
        c.fill = xl_fill('1E3A5F')
        c.alignment = xl_align('center', 'center')
        c.border = xl_border()

    indice_rows = (
        [('HU-01','Historia de Usuario: Consultar Servicios y Disponibilidad','Historia Usuario')] +
        [( f'HU-{i:02d}', f'Historia de Usuario: {HISTORIAS[i-1]["titulo"]}', 'Historia Usuario') for i in range(2,11)] +
        [('CU-01','Caso de Uso: Consultar Servicios y Disponibilidad','Caso de Uso')] +
        [( f'CU-{i:02d}', f'Caso de Uso: {CASOS_DE_USO[i-1]["nombre"]}', 'Caso de Uso') for i in range(2,11)] +
        [('Modelo BD','Tablas, columnas, tipos y FK del sistema','Referencia técnica')]
    )
    tipo_fills = {'Historia Usuario': 'FFF0F6', 'Caso de Uso': 'EFF6FF', 'Referencia técnica': 'F0FDF4'}
    for i, (hoja, contenido, tipo) in enumerate(indice_rows):
        r = i + 5
        ws_idx.row_dimensions[r].height = 18
        fill = xl_fill(tipo_fills.get(tipo, 'FFFFFF'))
        for ci, val in enumerate([hoja, contenido, tipo], 1):
            c = ws_idx.cell(row=r, column=ci, value=val)
            c.font = xl_font(bold=(ci==1), size=10,
                             color='C41E3A' if (ci==1 and 'HU' in hoja) else
                                   '1D4ED8' if (ci==1 and 'CU' in hoja) else
                                   '065F46' if ci==1 else '1E293B')
            c.fill = fill
            c.alignment = xl_align('left', 'center')
            c.border = xl_border()

    # ── Una hoja por Historia de Usuario ─────────────────────────────────────
    for hu in HISTORIAS:
        ws = wb.create_sheet(hu['id'])   # ej: "HU-01"
        set_cols(ws)
        write_hu_block(ws, hu, 1)

    # ── Una hoja por Caso de Uso ──────────────────────────────────────────────
    for cu in CASOS_DE_USO:
        ws = wb.create_sheet(cu['id'])   # ej: "CU-01"
        set_cols(ws)
        write_cu_block(ws, cu, 1)

    # ── Hoja 3: Modelo BD (10 tablas) ────────────────────────────────────────
    ws_bd = wb.create_sheet('Modelo BD')
    set_cols(ws_bd)
    ws_bd.column_dimensions['A'].width = 22
    ws_bd.column_dimensions['B'].width = 20
    ws_bd.column_dimensions['C'].width = 18
    ws_bd.column_dimensions['D'].width = 14
    ws_bd.column_dimensions['E'].width = 10
    ws_bd.column_dimensions['F'].width = 40

    ws_bd.row_dimensions[1].height = 30
    titulo = ws_bd.cell(row=1, column=1, value='MODELO DE BASE DE DATOS — 10 TABLAS ROSSMIX')
    titulo.font = xl_font(bold=True, size=14, color='C41E3A')
    titulo.fill = FILL_WHITE
    titulo.alignment = xl_align('left', 'center')
    ws_bd.merge_cells('A1:F1')

    hdrs_bd = ['Tabla', 'Columna', 'Tipo', 'Clave', 'Nullable', 'Descripción / FK']
    ws_bd.row_dimensions[2].height = 22
    for ci, h in enumerate(hdrs_bd, 1):
        c = ws_bd.cell(row=2, column=ci, value=h)
        c.font = xl_font(bold=True, size=10, color='FFFFFF')
        c.fill = xl_fill('1E3A5F')
        c.alignment = xl_align('center', 'center')
        c.border = xl_border()

    TABLAS_BD = [
        ('usuario','id','SERIAL','PK','NO','Clave primaria'),
        ('usuario','nombre','VARCHAR(100)','','NO','Nombre completo'),
        ('usuario','email','VARCHAR(150)','UK','NO','Email único'),
        ('usuario','telefono','VARCHAR(20)','','NO','Celular/WhatsApp'),
        ('usuario','password','VARCHAR(200)','','NO','Hash Werkzeug'),
        ('usuario','tipo_usuario','VARCHAR(20)','','NO','admin|cliente|especialista'),
        ('usuario','activo','BOOLEAN','','NO','Cuenta activa'),
        ('usuario','id_empleado','INTEGER','FK→empleados','SÍ','SET NULL al borrar empleado'),
        ('empleados','id_empleado','SERIAL','PK','NO','Clave primaria'),
        ('empleados','nombre','VARCHAR(100)','','NO','Nombre de la especialista'),
        ('empleados','especialidad','VARCHAR(100)','','SÍ','Área de especialización'),
        ('empleados','activo','BOOLEAN','','NO','Disponible para citas'),
        ('servicios','id_servicio','SERIAL','PK','NO','Clave primaria'),
        ('servicios','nombre_servicio','VARCHAR(100)','','NO','Nombre del servicio'),
        ('servicios','precio_total','NUMERIC(10,2)','','NO','Precio en COP'),
        ('servicios','duracion_minutos','INTEGER','','NO','Duración estimada'),
        ('servicios','activo','BOOLEAN','','NO','Publicado en catálogo'),
        ('empleado_servicios','id_empleado','INTEGER','FK→empleados (PK)','NO','Relación M:N'),
        ('empleado_servicios','id_servicio','INTEGER','FK→servicios (PK)','NO','Relación M:N'),
        ('horarios_empleados','id_horario','SERIAL','PK','NO','Clave primaria'),
        ('horarios_empleados','id_empleado','INTEGER','FK→empleados','NO','CASCADE al borrar'),
        ('horarios_empleados','dia_semana','INTEGER','','NO','0=Dom 1=Lun...6=Sáb'),
        ('horarios_empleados','hora_inicio','TIME','','NO','Inicio del turno'),
        ('horarios_empleados','hora_fin','TIME','','NO','Fin del turno'),
        ('citas','id_cita','SERIAL','PK','NO','Clave primaria'),
        ('citas','id_cliente','INTEGER','FK→usuario','NO','CASCADE al borrar usuario'),
        ('citas','id_empleado','INTEGER','FK→empleados','SÍ','SET NULL al borrar empleado'),
        ('citas','id_servicio','INTEGER','FK→servicios','NO','RESTRICT al borrar servicio'),
        ('citas','estado','estado_cita_enum','','NO','pendiente_pago|confirmada|en_atencion|completada|cancelada|no_asistio'),
        ('citas','monto_abono','NUMERIC(10,2)','','SÍ','Abono mínimo $5.000 COP'),
        ('citas','saldo_pendiente','NUMERIC(10,2)','','SÍ','monto_total - monto_abono'),
        ('citas','codigo_reserva','VARCHAR(20)','UK','SÍ','Generado por CitaService (8 chars)'),
        ('citas','token_gestion','VARCHAR(32)','UK','SÍ','secrets.token_urlsafe(24)'),
        ('pagos','id_pago','SERIAL','PK','NO','Clave primaria'),
        ('pagos','id_cita','INTEGER','FK→citas UK','NO','1 pago por cita — CASCADE'),
        ('pagos','monto','NUMERIC(10,2)','','NO','Monto cobrado'),
        ('pagos','metodo_pago','metodo_pago_enum','','NO','efectivo|tarjeta|transferencia|nequi|daviplata'),
        ('pagos','estado_pago','VARCHAR(20)','','NO','pendiente|completado|reembolsado'),
        ('notificaciones','id','SERIAL','PK','NO','Clave primaria'),
        ('notificaciones','id_usuario','INTEGER','FK→usuario','NO','CASCADE al borrar usuario'),
        ('notificaciones','titulo','VARCHAR(200)','','NO','Asunto'),
        ('notificaciones','leido','BOOLEAN','','NO','Estado de lectura'),
        ('notificaciones','target','VARCHAR(300)','','SÍ','URL de destino al hacer clic'),
        ('auditoria_usuarios','id','SERIAL','PK','NO','Clave primaria'),
        ('auditoria_usuarios','id_usuario','INTEGER','FK→usuario','SÍ','SET NULL — usuario afectado'),
        ('auditoria_usuarios','id_actor','INTEGER','FK→usuario','SÍ','SET NULL — admin que actuó'),
        ('auditoria_usuarios','accion','VARCHAR(50)','','NO','login|logout|editar|eliminar|desactivar'),
        ('auditoria_usuarios','ip_address','VARCHAR(45)','','SÍ','IPv4 o IPv6'),
        ('auditoria_usuarios','fecha','TIMESTAMP','','NO','DEFAULT CURRENT_TIMESTAMP'),
        ('configuraciones','id','SERIAL','PK','NO','Clave primaria'),
        ('configuraciones','clave','VARCHAR(120)','UK','NO','Identificador del parámetro'),
        ('configuraciones','valor','TEXT','','NO','Valor en texto plano'),
        ('configuraciones','creado_por','INTEGER','FK→usuario','SÍ','SET NULL — admin creador'),
        ('configuraciones','modificado_por','INTEGER','FK→usuario','SÍ','SET NULL — último editor'),
    ]

    tabla_colors = {
        'usuario':'FFF0F6', 'empleados':'F0FDF4', 'servicios':'F0FDF4',
        'empleado_servicios':'EFF6FF', 'horarios_empleados':'EFF6FF',
        'citas':'FFF7ED', 'pagos':'FFFBEB',
        'notificaciones':'FDF4FF', 'auditoria_usuarios':'F8FAFC',
        'configuraciones':'F0F9FF',
    }
    for i, fila in enumerate(TABLAS_BD):
        r = i + 3
        ws_bd.row_dimensions[r].height = 18
        fill = xl_fill(tabla_colors.get(fila[0], 'FFFFFF'))
        for ci, val in enumerate(fila, 1):
            c = ws_bd.cell(row=r, column=ci, value=val)
            if ci == 1:
                c.font = xl_font(bold=True, size=9, color='065F46')
            elif ci == 4 and val:
                color = 'C41E3A' if 'PK' in val else '1D4ED8'
                c.font = xl_font(bold=True, size=9, color=color)
            else:
                c.font = xl_font(size=9)
            c.fill = fill
            c.alignment = xl_align('left', 'center', wrap=True)
            c.border = xl_border()

    ws_bd.freeze_panes = 'A3'

    os.makedirs('docs', exist_ok=True)
    out = 'docs/Rossmix_CasosDeUso_v2.xlsx'
    wb.save(out)
    print(f'[OK] Excel: {out}')


# ── Ejecutar ──────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    generar_word()
    generar_excel()
    print('\n✓ Documentos generados en docs/')
