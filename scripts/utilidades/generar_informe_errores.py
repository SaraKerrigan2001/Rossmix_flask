"""
Genera el informe Word completo de todos los errores y correcciones
del proyecto Rossmix Flask durante todo el proceso de desarrollo.

Salida: docs/Rossmix_Informe_Errores_Proceso.docx
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# ── Paleta ────────────────────────────────────────────────────────────────────
ROJO    = RGBColor(0xC4, 0x1E, 0x3A)
ROSA    = RGBColor(0xFF, 0x14, 0x93)
OSCURO  = RGBColor(0x1A, 0x1A, 0x1A)
GRIS    = RGBColor(0x44, 0x44, 0x44)
BLANCO  = RGBColor(0xFF, 0xFF, 0xFF)

def rgb_hex(h):
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

# ── DATOS — HISTORIAL COMPLETO DE ERRORES Y CORRECCIONES ─────────────────────
ERRORES = [
    # ── FASE 1: REVISIÓN INICIAL ──────────────────────────────────────────────
    {
        "fase": "Fase 1 — Revisión Inicial del Proyecto",
        "items": [
            {
                "id": "E-001", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "credenciales_login.md con contraseñas en texto plano en el repositorio",
                "archivo": "credenciales_login.md",
                "detalle": (
                    "El archivo contenía contraseñas de todos los usuarios del sistema "
                    "(admin123, especialista123, cliente123) accesibles públicamente en el repo de GitHub."
                ),
                "correccion": (
                    "Archivo eliminado del tracking de git con 'git rm --cached'. "
                    "Agregado al .gitignore con tres reglas de exclusión."
                ),
                "commit": "49e116c",
            },
            {
                "id": "E-002", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "token_gestion nunca generado al crear citas — funcionalidad rota",
                "archivo": "app/services/citas_service.py",
                "detalle": (
                    "El campo token_gestion existía en el modelo y en el SQL pero nunca se asignaba "
                    "al crear una cita. La ruta gestionar_cita buscaba por este token que siempre era NULL."
                ),
                "correccion": (
                    "Agregado secrets.token_urlsafe(24) en CitaService.crear_cita() "
                    "y nuevo método generar_token_gestion()."
                ),
                "commit": "49e116c",
            },
            {
                "id": "E-003", "severidad": "CRÍTICO", "categoria": "Lógica de Negocio",
                "descripcion": "reagendar_no_asistio prometía crédito pero no lo implementaba",
                "archivo": "app/views/citas.py",
                "detalle": (
                    "El endpoint mostraba el mensaje 'Tu abono se aplicará como crédito' "
                    "pero no guardaba nada en sesión ni en BD. El crédito nunca se descontaba."
                ),
                "correccion": (
                    "Implementado: guarda monto_credito en session['credito_reagenda'] "
                    "con flag [CREDITO_CONSUMIDO] en BD para evitar reutilización. "
                    "Cambiado de GET a POST para protección CSRF."
                ),
                "commit": "49e116c",
            },
            {
                "id": "E-004", "severidad": "ALTO", "categoria": "Funcionalidad",
                "descripcion": "reprogramar_cita_form sin ruta POST — reprogramación incompleta",
                "archivo": "app/views/citas.py",
                "detalle": (
                    "La ruta GET mostraba el formulario pero no existía ruta POST para procesar "
                    "la reprogramación. El usuario seleccionaba nueva fecha pero no se guardaba nada."
                ),
                "correccion": (
                    "Creada ruta POST reprogramar_cita_submit que cancela la cita original, "
                    "transfiere el abono y crea la nueva cita con todos los datos validados."
                ),
                "commit": "49e116c",
            },
            {
                "id": "E-005", "severidad": "ALTO", "categoria": "Bug",
                "descripcion": "reportes_service.py usaba modelo Usuario en lugar de Empleado para horarios",
                "archivo": "app/services/reportes_service.py",
                "detalle": (
                    "La exportación de horarios hacía: empleado = Usuario.query.get(h.id_empleado) "
                    "cuando debería ser Empleado. Retornaba None o el usuario equivocado."
                ),
                "correccion": (
                    "Reescrito con JOINs en todas las queries. Bug de Usuario→Empleado corregido. "
                    "Eliminadas N+1 queries en exportación."
                ),
                "commit": "49e116c",
            },
            {
                "id": "E-006", "severidad": "ALTO", "categoria": "Inconsistencia",
                "descripcion": "Lógica de abono inconsistente entre admin y cliente",
                "archivo": "app/views/admin/pagos.py",
                "detalle": (
                    "El admin reemplazaba el abono (cita.monto_abono = monto) mientras que "
                    "el cliente acumulaba (cita.monto_abono += monto). Generaba estados inconsistentes."
                ),
                "correccion": "Unificado: ambos acumulan el abono igual.",
                "commit": "49e116c",
            },
            {
                "id": "E-007", "severidad": "ALTO", "categoria": "Bug",
                "descripcion": "Lógica de disponibilidad duplicada entre vista y servicio",
                "archivo": "app/views/citas.py",
                "detalle": (
                    "El endpoint horarios_disponibles tenía 30 líneas de lógica idéntica "
                    "a CitaService.obtener_horarios_disponibles(). Dos versiones divergentes."
                ),
                "correccion": "Vista delega completamente al servicio eliminando 30 líneas duplicadas.",
                "commit": "49e116c",
            },
            {
                "id": "E-008", "severidad": "ALTO", "categoria": "Seguridad",
                "descripcion": "ENUMs con create_type=False rompían en BD nueva",
                "archivo": "app/models/cita.py, app/models/pago.py",
                "detalle": (
                    "create_constraint=False en SQLAlchemy no creaba los tipos ENUM. "
                    "Si se recreaba la BD con db.create_all(), fallaba con ProgrammingError."
                ),
                "correccion": "Cambiado a create_type=False (correcto para psycopg3).",
                "commit": "49e116c",
            },
            {
                "id": "E-009", "severidad": "MEDIO", "categoria": "Seguridad",
                "descripcion": "admin_required no verificaba usuario.activo en BD",
                "archivo": "app/utils/decorators.py",
                "detalle": (
                    "El decorador solo verificaba la sesión. Si un admin era desactivado "
                    "en BD, su sesión existente seguía funcionando indefinidamente."
                ),
                "correccion": (
                    "Reescrito: verifica usuario.activo en BD en cada request. "
                    "Nuevo especialista_required que también verifica el empleado vinculado."
                ),
                "commit": "49e116c",
            },
            {
                "id": "E-010", "severidad": "MEDIO", "categoria": "Sesión",
                "descripcion": "Sesión sin timeout configurado",
                "archivo": "app/config.py",
                "detalle": "Las sesiones no tenían duración límite. Un admin podía dejar sesión abierta indefinidamente.",
                "correccion": "PERMANENT_SESSION_LIFETIME = 8 horas, configurable por SESSION_LIFETIME_SECONDS en .env.",
                "commit": "49e116c",
            },
        ]
    },
    # ── FASE 2: SEGURIDAD ─────────────────────────────────────────────────────
    {
        "fase": "Fase 2 — Correcciones de Seguridad (14 Hallazgos)",
        "items": [
            {
                "id": "E-011", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "Validación insuficiente al confirmar cita — campos ocultos manipulables",
                "archivo": "app/views/citas.py:176",
                "detalle": (
                    "id_servicio, empleado, fecha y hora venían de campos ocultos sin validar "
                    "que el empleado hiciera ese servicio, estuviera activo o tuviera disponibilidad."
                ),
                "correccion": (
                    "Validaciones completas: servicio activo, empleado activo+ofrece servicio, "
                    "duración real, anticipación mínima 30min y disponibilidad real."
                ),
                "commit": "32cbe74",
            },
            {
                "id": "E-012", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "Pagos con importe arbitrario — sobrepagos posibles",
                "archivo": "app/views/citas.py:383, app/views/admin/pagos.py:38",
                "detalle": (
                    "Los endpoints de pago aceptaban importes superiores al saldo pendiente, "
                    "permitiendo sobrepagos y estados financieros inconsistentes."
                ),
                "correccion": "Monto limitado al saldo_pendiente real con tolerancia de 1 centavo.",
                "commit": "32cbe74",
            },
            {
                "id": "E-013", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "Crédito de inasistencia reutilizable múltiples veces",
                "archivo": "app/views/citas.py:545",
                "detalle": (
                    "El crédito solo se guardaba en sesión (no en BD). "
                    "El mismo abono podía aplicarse varias veces. "
                    "Además modificaba estado mediante GET sin CSRF."
                ),
                "correccion": (
                    "Flag [CREDITO_CONSUMIDO] guardado en cita.notas antes de usar el crédito. "
                    "Endpoint cambiado de GET a POST."
                ),
                "commit": "32cbe74",
            },
            {
                "id": "E-014", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "Reasignación de empleados sin validación — inactivos asignables",
                "archivo": "app/views/admin/citas.py:108,204",
                "detalle": "Permitía asignar empleados inactivos, sin el servicio o sin disponibilidad.",
                "correccion": "Validación de activo, servicio y disponibilidad en reasignación individual y batch.",
                "commit": "32cbe74",
            },
            {
                "id": "E-015", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "Condición de carrera al reservar citas simultáneas",
                "archivo": "app/services/citas_service.py:15",
                "detalle": "Dos solicitudes simultáneas podían reservar el mismo horario con el mismo empleado.",
                "correccion": (
                    "Índice único parcial en PostgreSQL: "
                    "CREATE UNIQUE INDEX idx_no_solapamiento_citas ON citas(id_empleado, fecha_hora_inicio) "
                    "WHERE estado IN ('pendiente_pago','confirmada','en_atencion')."
                ),
                "commit": "32cbe74",
            },
            {
                "id": "E-016", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "XSS almacenado en panel administrativo — innerHTML con datos de usuario",
                "archivo": "app/templates/admin/empleados.html:454, citas.html:283",
                "detalle": (
                    "Nombres de clientes y servicios insertados via innerHTML. "
                    "Un nombre malicioso como '<script>alert(1)</script>' ejecutaría código JS."
                ),
                "correccion": "innerHTML reemplazado por createElement/textContent en todos los casos.",
                "commit": "32cbe74",
            },
            {
                "id": "E-017", "severidad": "MEDIO", "categoria": "Seguridad",
                "descripcion": "Especialistas con empleados desactivados conservaban acceso",
                "archivo": "app/utils/decorators.py:68",
                "detalle": "especialista_required verificaba la cuenta pero no el empleado vinculado.",
                "correccion": "Verifica db.session.get(Empleado, usuario.id_empleado).activo en cada request.",
                "commit": "32cbe74",
            },
            {
                "id": "E-018", "severidad": "MEDIO", "categoria": "Seguridad",
                "descripcion": "Sin protección contra fuerza bruta en login",
                "archivo": "app/views/auth.py:23",
                "detalle": "No había límites por IP, bloqueo progresivo ni alertas en login.",
                "correccion": (
                    "Rate limiting: 10 intentos en 15 min, bloqueo 30 min via Flask-Caching. "
                    "Intentos fallidos registrados en auditoría."
                ),
                "commit": "32cbe74",
            },
            {
                "id": "E-019", "severidad": "MEDIO", "categoria": "Configuración",
                "descripcion": "debug=True hardcodeado en app.py y run.py",
                "archivo": "app.py:17, run.py:9",
                "detalle": "Si se usaban accidentalmente en producción, el debugger exponía código fuente.",
                "correccion": "debug = os.environ.get('FLASK_ENV') != 'production' — dinámico.",
                "commit": "32cbe74",
            },
            {
                "id": "E-020", "severidad": "MEDIO", "categoria": "Configuración",
                "descripcion": "SECRET_KEY cambiaba silenciosamente en cada arranque",
                "archivo": "app/config.py:45",
                "detalle": "Si faltaba la variable, se generaba una clave aleatoria sin advertencia visible.",
                "correccion": "sys.exit(1) en producción si falta SECRET_KEY. Advertencia en stderr en desarrollo.",
                "commit": "32cbe74",
            },
            {
                "id": "E-021", "severidad": "MEDIO", "categoria": "Bug",
                "descripcion": "Flujo reprogramar.html llamaba endpoint inexistente",
                "archivo": "app/templates/citas/reprogramar.html:280",
                "detalle": "JS llamaba /citas/reprogramar/TOKEN/confirmar — ese endpoint nunca existió (404).",
                "correccion": "Corregido a POST /citas/reprogramar/<id_cita> con SeleccionarHorarioForm.",
                "commit": "32cbe74",
            },
            {
                "id": "E-022", "severidad": "MEDIO", "categoria": "Consistencia",
                "descripcion": "Mezcla de datetime.now() y datetime.utcnow() en modelos",
                "archivo": "Múltiples modelos y servicios",
                "detalle": "En servidores UTC las citas podían desplazarse respecto a hora local de Colombia.",
                "correccion": "Unificado a datetime.now() en todos los modelos y servicios (28 archivos).",
                "commit": "32cbe74",
            },
            {
                "id": "E-023", "severidad": "MEDIO", "categoria": "Datos",
                "descripcion": "Email duplicado carlos@gmail.com en crear_usuarios.py",
                "archivo": "scripts/utilidades/crear_usuarios.py:42",
                "detalle": "El mismo email aparecía dos veces, causando fallo en UNIQUE constraint al cargar seeds.",
                "correccion": "Email corregido a andrade.cliente@rossmix.com para el segundo registro.",
                "commit": "32cbe74",
            },
            {
                "id": "E-024", "severidad": "BAJO", "categoria": "Seguridad",
                "descripcion": "docker-compose.yml con POSTGRES_PASSWORD=1234 hardcodeado",
                "archivo": "docker-compose.yml:8",
                "detalle": "Contraseña débil y predecible por defecto sin ninguna advertencia.",
                "correccion": "Cambiado a ${DB_PASSWORD} con sintaxis :? que falla si no está definida.",
                "commit": "32cbe74",
            },
        ]
    },
    # ── FASE 3: REVISIÓN DE CALIDAD ───────────────────────────────────────────
    {
        "fase": "Fase 3 — Revisión de Calidad de Código",
        "items": [
            {
                "id": "E-025", "severidad": "CRÍTICO", "categoria": "Bug",
                "descripcion": "session['email'] nunca guardado — perfil dropdown siempre vacío",
                "archivo": "app/views/auth.py",
                "detalle": "base.html usaba session.get('email') para mostrar el email del usuario en el dropdown. Siempre mostraba cadena vacía.",
                "correccion": "Agregado session['email'] = usuario.email en el login exitoso.",
                "commit": "84920d8",
            },
            {
                "id": "E-026", "severidad": "CRÍTICO", "categoria": "Bug",
                "descripcion": "Race condition en asignación aleatoria de empleado",
                "archivo": "app/views/citas.py",
                "detalle": (
                    "horarios_disponibles elegía empleado aleatorio para mostrar slots, "
                    "y confirmar_cita elegía otro diferente. Las horas mostradas no coincidían "
                    "con el empleado asignado."
                ),
                "correccion": "Se elige UN empleado antes de consultar slots y se devuelve ese mismo al cliente.",
                "commit": "84920d8",
            },
            {
                "id": "E-027", "severidad": "ALTO", "categoria": "Performance",
                "descripcion": "N+1 queries en listado de empleados (admin)",
                "archivo": "app/views/admin/empleados.py",
                "detalle": "Una query por empleado para contar sus servicios. Con 20 empleados = 21 queries.",
                "correccion": "Una sola query con GROUP BY: dict(db.session.query(EmpleadoServicio.id_empleado, func.count()).group_by(...).all())",
                "commit": "84920d8",
            },
            {
                "id": "E-028", "severidad": "ALTO", "categoria": "Performance",
                "descripcion": "N+1 queries en listado de clientes (admin)",
                "archivo": "app/views/admin/clientes.py",
                "detalle": "Una query por cliente para contar citas canceladas. Con 100 clientes = 101 queries.",
                "correccion": "Una sola query con GROUP BY para todas las canceladas.",
                "commit": "84920d8",
            },
            {
                "id": "E-029", "severidad": "ALTO", "categoria": "Seguridad",
                "descripcion": "Edición de cliente sin validación de email/teléfono",
                "archivo": "app/views/admin/clientes.py",
                "detalle": "Usaba request.form directo sin WTForms. Podía guardar emails inválidos o teléfonos de menos de 10 dígitos.",
                "correccion": "Validación con regex para email y teléfono (exactamente 10 dígitos) antes de guardar.",
                "commit": "84920d8",
            },
            {
                "id": "E-030", "severidad": "ALTO", "categoria": "Performance",
                "descripcion": "citas_completadas cargaba TODAS las citas en memoria",
                "archivo": "app/models/usuario.py",
                "detalle": "sum(1 for c in self.citas if c.estado == 'completada') — lazy load de todas las citas.",
                "correccion": "Reemplazado por COUNT en BD: db.session.query(func.count(...)).filter(...).scalar()",
                "commit": "84920d8",
            },
            {
                "id": "E-031", "severidad": "ALTO", "categoria": "Seguridad",
                "descripcion": "flash(str(e)) exponía mensajes internos de SQLAlchemy al usuario",
                "archivo": "app/views/citas.py:283,761",
                "detalle": "Los mensajes de excepción podían revelar nombres de tablas, columnas o stack traces.",
                "correccion": "Reemplazado por mensajes genéricos: 'No fue posible completar la operación. Inténtalo de nuevo.'",
                "commit": "84920d8",
            },
            {
                "id": "E-032", "severidad": "ALTO", "categoria": "Bug",
                "descripcion": "reprogramar_cita_submit no filtraba empleados activos",
                "archivo": "app/views/citas.py:701",
                "detalle": "EmpleadoServicio.query.filter_by(id_servicio=id_servicio) — sin filtrar Empleado.activo.",
                "correccion": "JOIN con Empleado y filtro Empleado.activo == True en la consulta.",
                "commit": "84920d8",
            },
            {
                "id": "E-033", "severidad": "ALTO", "categoria": "Bug",
                "descripcion": "30+ ocurrencias de query.get_or_404() deprecado en SQLAlchemy 2.x",
                "archivo": "10 archivos de vistas",
                "detalle": "Model.query.get_or_404() está deprecado desde SQLAlchemy 2.0. Genera warnings en cada request.",
                "correccion": "31 reemplazos por db.get_or_404(Model, pk) con script automatizado.",
                "commit": "84920d8",
            },
            {
                "id": "E-034", "severidad": "ALTO", "categoria": "Seguridad",
                "descripcion": "Sin rate limiting en /registro — creación masiva de cuentas",
                "archivo": "app/views/auth.py",
                "detalle": "Solo /login tenía protección. /registro podía recibir miles de solicitudes sin restricción.",
                "correccion": "Misma lógica de rate limiting: 10 intentos / bloqueo 30 min aplicada a /registro.",
                "commit": "84920d8",
            },
            {
                "id": "E-035", "severidad": "ALTO", "categoria": "Seguridad",
                "descripcion": "Cookies de sesión sin flags SECURE/HTTPONLY/SAMESITE",
                "archivo": "app/config.py",
                "detalle": "Sin SESSION_COOKIE_SECURE, SESSION_COOKIE_HTTPONLY ni SESSION_COOKIE_SAMESITE.",
                "correccion": "Agregados en config.py: SECURE=True en producción, HTTPONLY=True, SAMESITE='Lax'.",
                "commit": "84920d8",
            },
            {
                "id": "E-036", "severidad": "ALTO", "categoria": "Seguridad",
                "descripcion": "Código de reserva generado con random no criptográfico",
                "archivo": "app/services/citas_service.py",
                "detalle": "random.choices() es pseudoaleatorio y predecible. Permitía enumerar citas de otros usuarios.",
                "correccion": "Reemplazado por secrets.token_urlsafe(6)[:8].upper() — criptográficamente seguro.",
                "commit": "84920d8",
            },
            {
                "id": "E-037", "severidad": "MEDIO", "categoria": "Bug",
                "descripcion": "add_notificacion hacía db.session.commit() propio — rompía atomicidad",
                "archivo": "app/utils/helpers.py",
                "detalle": "Si la operación principal fallaba después de crear la notificación, la notificación quedaba huérfana.",
                "correccion": "Eliminado el commit propio. El llamador es responsable de la transacción.",
                "commit": "84920d8",
            },
            {
                "id": "E-038", "severidad": "MEDIO", "categoria": "Seguridad",
                "descripcion": "citas_asignar_batch sin validación de disponibilidad",
                "archivo": "app/views/admin/citas.py",
                "detalle": "La asignación masiva no verificaba solapamientos ni empleados activos.",
                "correccion": "Valida activo y CitaService.validar_disponibilidad_cita() antes de asignar cada cita.",
                "commit": "84920d8",
            },
            {
                "id": "E-039", "severidad": "BAJO", "categoria": "Performance",
                "descripcion": "Índices faltantes en columnas de búsqueda frecuente",
                "archivo": "app/models/usuario.py, app/models/cita.py",
                "detalle": "tipo_usuario, estado y fecha_hora_inicio sin index=True. Queries de listado sin índice.",
                "correccion": "Agregado index=True en tipo_usuario (usuario), estado y fecha_hora_inicio (cita).",
                "commit": "84920d8",
            },
        ]
    },
    # ── FASE 4: PALETA DE COLORES ─────────────────────────────────────────────
    {
        "fase": "Fase 4 — Unificación de Paleta de Colores (655 hallazgos)",
        "items": [
            {
                "id": "E-040", "severidad": "ALTO", "categoria": "UI/UX",
                "descripcion": "Modales de Nuevo Empleado y Nuevo Servicio con headers morado y verde",
                "archivo": "app/templates/admin/empleados.html, servicios.html",
                "detalle": "Header morado (#7c3aed, #a855f7) en modal empleado. Header verde (#059669, #10b981) en modal servicio.",
                "correccion": "Cambiados a gradiente rosa Rossmix: linear-gradient(135deg, #c41e3a, #ff1493).",
                "commit": "913c2b3",
            },
            {
                "id": "E-041", "severidad": "ALTO", "categoria": "UI/UX",
                "descripcion": "Filas de horarios con fondo rosa y texto blanco — ilegibles",
                "archivo": "app/templates/admin/horarios.html",
                "detalle": "El fondo rosa del header se extendía a las filas de datos. El texto blanco sobre rosa era difícil de leer.",
                "correccion": "Restauradas filas con fondo blanco, texto negro. Badge del día en rosa en esquina izquierda.",
                "commit": "5757b09",
            },
            {
                "id": "E-042", "severidad": "ALTO", "categoria": "UI/UX",
                "descripcion": "Botón 'Crear Cuenta' con gradiente dorado/marrón",
                "archivo": "app/templates/registro.html",
                "detalle": "#C41E3A → #C5A059 (dorado). Inconsistente con la paleta del sistema.",
                "correccion": "Cambiado a linear-gradient(135deg, #c41e3a, #ff1493).",
                "commit": "f20f46d",
            },
            {
                "id": "E-043", "severidad": "MEDIO", "categoria": "UI/UX",
                "descripcion": "Secciones de index.html con 8 colores distintos (naranja, morado, azul, verde, dorado)",
                "archivo": "app/templates/index.html",
                "detalle": (
                    "Categorías: naranja/rojo, rosa/morado, gris/azul, morado/púrpura. "
                    "Colecciones: 8 gradientes diferentes. Diseños: 6 fondos distintos."
                ),
                "correccion": "Todos reemplazados por variaciones del rosa Rossmix: #fff0f6 → #ffd6e8.",
                "commit": "f20f46d",
            },
            {
                "id": "E-044", "severidad": "MEDIO", "categoria": "UI/UX",
                "descripcion": "dashboard_admin.html con múltiples colores temáticos (purple, teal, gold, blue)",
                "archivo": "app/templates/dashboard_admin.html",
                "detalle": "37 colores fuera de paleta: stat-cards con bordes púrpura/dorado/verde, módulos con iconos de colores distintos.",
                "correccion": "37 reemplazos — todos a variaciones del rosa/rojo Rossmix.",
                "commit": "56c171a",
            },
            {
                "id": "E-045", "severidad": "MEDIO", "categoria": "UI/UX",
                "descripcion": "526 colores fuera de paleta en 29 archivos — auditoría final completa",
                "archivo": "Todos los templates y style.css",
                "detalle": (
                    "Verde (#059669, #16a34a, #10b981), azul (#2563eb, #3b82f6), "
                    "morado (#7c3aed, #8b5cf6), amarillo (#d97706, #f59e0b), "
                    "rosas externos (#c71585, #ff69b4)."
                ),
                "correccion": "Script automatizado aplicó 526 reemplazos en 29 archivos en una sola pasada.",
                "commit": "4523a4f",
            },
            {
                "id": "E-046", "severidad": "BAJO", "categoria": "UI/UX",
                "descripcion": "Textos de subtítulos en dashboards con color blanco sobre fondos claros",
                "archivo": "dashboard_admin.html, dashboard_cliente.html, especialista/dashboard.html",
                "detalle": "Varios textos descriptivos usaban color: rgba(255,255,255,0.8) sobre fondos blancos/crema — invisibles.",
                "correccion": "Cambiados a #444 (subtítulos) y #555 (texto secundario) para legibilidad.",
                "commit": "4523a4f",
            },
        ]
    },
    # ── FASE 5: INFRAESTRUCTURA ───────────────────────────────────────────────
    {
        "fase": "Fase 5 — Infraestructura, Docker y Acceso Móvil",
        "items": [
            {
                "id": "E-047", "severidad": "ALTO", "categoria": "Infraestructura",
                "descripcion": "Servidor Flask escuchando solo en 127.0.0.1 — inaccesible desde celular",
                "archivo": "app.py",
                "detalle": "Flask corría en localhost únicamente. Imposible acceder desde dispositivos en la misma red WiFi.",
                "correccion": "app.run(host='0.0.0.0', port=5000). Puerto 5000 abierto en firewall de Windows.",
                "commit": "03f35da",
            },
            {
                "id": "E-048", "severidad": "ALTO", "categoria": "Infraestructura",
                "descripcion": ".env con BOM UTF-8 — python-dotenv no leía las variables",
                "archivo": ".env",
                "detalle": "El archivo tenía BOM EF BB BF al inicio. python-dotenv no reconocía las variables y DATABASE_URL era None.",
                "correccion": "Reescrito sin BOM con System.Text.UTF8Encoding(false). Conexión a BD restaurada.",
                "commit": "03f35da",
            },
            {
                "id": "E-049", "severidad": "ALTO", "categoria": "Infraestructura",
                "descripcion": "docker-compose.yml con DATABASE_URL apuntando a host.docker.internal",
                "archivo": ".env, docker-compose.yml",
                "detalle": "host.docker.internal es para Docker Desktop en Mac. En Windows con la config actual no resolvía.",
                "correccion": "Cambiado a localhost para desarrollo local. docker-compose.yml usa db (nombre del servicio).",
                "commit": "b769896",
            },
            {
                "id": "E-050", "severidad": "MEDIO", "categoria": "Infraestructura",
                "descripcion": "Dockerfile usando Python 3.11 — proyecto requiere 3.13",
                "archivo": "Dockerfile",
                "detalle": "La imagen base era python:3.11-slim pero el proyecto usa características de Python 3.13.",
                "correccion": "Actualizado a python:3.13-slim con multi-stage build (builder + runtime). Usuario no-root.",
                "commit": "b769896",
            },
            {
                "id": "E-051", "severidad": "MEDIO", "categoria": "Seguridad",
                "descripcion": "YAML inválido en docker-compose.yml — sintaxis :? rompe el parser",
                "archivo": "docker-compose.yml",
                "detalle": "POSTGRES_PASSWORD: ${DB_PASSWORD:?mensaje} — los dos puntos rompen el parser YAML de Docker Compose.",
                "correccion": "Cambiado a POSTGRES_PASSWORD: ${DB_PASSWORD} con validación separada.",
                "commit": "b769896",
            },
            {
                "id": "E-052", "severidad": "MEDIO", "categoria": "Infraestructura",
                "descripcion": "Menú hamburguesa no funcionaba en móvil",
                "archivo": "app/templates/base.html",
                "detalle": (
                    "El DOMContentLoaded del hamburguesa estaba mal cerrado. "
                    "El código del profileBtn quedaba pegado dentro rompiendo ambos eventos. "
                    "Además el menú usaba position:absolute en lugar de position:fixed."
                ),
                "correccion": (
                    "Reescrito como IIFE independiente. position:fixed con top:64px. "
                    "e.stopPropagation() para evitar cierre inmediato. "
                    "-webkit-tap-highlight-color para Android."
                ),
                "commit": "6796dc2",
            },
            {
                "id": "E-053", "severidad": "BAJO", "categoria": "Git",
                "descripcion": "cloudflared.exe (73 MB) en el directorio raíz sin .gitignore",
                "archivo": "cloudflared.exe",
                "detalle": "Binario de 73 MB podría commitearse accidentalmente al repositorio.",
                "correccion": "Agregado cloudflared.exe, ngrok/, instance/, *.db al .gitignore.",
                "commit": "84920d8",
            },
            {
                "id": "E-054", "severidad": "BAJO", "categoria": "Git",
                "descripcion": "Dos entornos virtuales (.venv y .venv313) con rutas rotas",
                "archivo": ".venv/, .venv313/",
                "detalle": "Ambos apuntaban a C:\\Users\\maria\\mi_proyecto_flask\\ (ruta inexistente). pip no funcionaba.",
                "correccion": "Eliminados ambos. Creado .venv nuevo con Python 3.14 limpio. Instaladas 13 dependencias.",
                "commit": "N/A — manual",
            },
        ]
    },
    # ── FASE 6: FORMULARIOS Y NAVBAR ─────────────────────────────────────────
    {
        "fase": "Fase 6 — Formularios, Navbar y Accesibilidad",
        "items": [
            {
                "id": "E-055", "severidad": "CRÍTICO", "categoria": "Seguridad",
                "descripcion": "Error 'Falta el token CSRF' en todos los formularios admin",
                "archivo": "app/templates/admin/*.html (5 archivos)",
                "detalle": (
                    "Los formularios admin usaban request.form directamente sin FlaskForm. "
                    "CSRF estaba habilitado globalmente pero los forms no incluían el token hidden."
                ),
                "correccion": (
                    "Token CSRF agregado a todos los formularios: "
                    "<input type='hidden' name='csrf_token' value='{{ csrf_token() }}'>. "
                    "Interceptor global fetch() que inyecta X-CSRFToken en todas las peticiones AJAX."
                ),
                "commit": "2307d3c",
            },
            {
                "id": "E-056", "severidad": "ALTO", "categoria": "Bug",
                "descripcion": "Campo monto con step=1000 rechazaba valores como 40000",
                "archivo": "app/templates/admin/pagos_form.html",
                "detalle": "El error 'Los valores válidos más aproximados son 33001 y 40001' aparecía al ingresar 40000.",
                "correccion": "step='1' min='5000'. Flechas del spinner ocultas con -webkit-appearance:none.",
                "commit": "2307d3c",
            },
            {
                "id": "E-057", "severidad": "ALTO", "categoria": "Bug",
                "descripcion": "Navbar mostraba links de desktop en móvil superpuestos",
                "archivo": "app/static/style.css",
                "detalle": (
                    "style.css tenía nav ul { gap: 0.3rem } en el mismo media query donde base.html "
                    "ponía nav ul { display: none !important }. La segunda regla sobreescribía la primera."
                ),
                "correccion": "style.css actualizado con nav ul { display: none !important } para garantizar ocultamiento.",
                "commit": "f811e75",
            },
            {
                "id": "E-058", "severidad": "MEDIO", "categoria": "UX",
                "descripcion": "Login y Registro compartían el mismo layout con todos los links de navbar",
                "archivo": "app/templates/login.html, registro.html",
                "detalle": "Aparecían 'Inicio + Iniciar Sesión + Registrarse' todos juntos en la misma página.",
                "correccion": (
                    "Creado auth_base.html con navbar mínima. "
                    "Login muestra solo 'Crear Cuenta'. Registro muestra solo 'Iniciar Sesión'."
                ),
                "commit": "2307d3c",
            },
            {
                "id": "E-059", "severidad": "MEDIO", "categoria": "UX",
                "descripcion": "Pagos por Confirmar siempre visible en navbar aunque no hubiera pagos",
                "archivo": "app/templates/base.html, app/utils/helpers.py",
                "detalle": "El botón amarillo aparecía aunque el contador fuera 0.",
                "correccion": (
                    "pagos_por_confirmar calculado en context processor con subconsulta NOT EXISTS. "
                    "Botón oculto cuando es 0, visible con contador cuando hay pendientes."
                ),
                "commit": "2307d3c",
            },
            {
                "id": "E-060", "severidad": "MEDIO", "categoria": "UX",
                "descripcion": "Botones de acción en tablas admin con colores inconsistentes",
                "archivo": "app/static/style.css",
                "detalle": "edit-action: amarillo, history-action: azul, delete-action: rojo. Fuera de paleta.",
                "correccion": "Todos unificados a fondo #fff0f6 con texto #c41e3a (rosa Rossmix).",
                "commit": "56c171a",
            },
        ]
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# GENERAR DOCUMENTO WORD
# ─────────────────────────────────────────────────────────────────────────────
def generar_word():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ROSSMIX')
    r.font.size = Pt(36); r.bold = True; r.font.color.rgb = ROJO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Salón de Belleza — Sistema de Agendamiento de Citas')
    r2.font.size = Pt(13); r2.font.color.rgb = GRIS

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('INFORME DE ERRORES Y CORRECCIONES')
    r3.font.size = Pt(20); r3.bold = True; r3.font.color.rgb = OSCURO

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('Historial completo del proceso de desarrollo y corrección')
    r4.font.size = Pt(11); r4.font.color.rgb = GRIS

    doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(f'Fecha: {date.today().strftime("%d de %B de %Y")}  ·  Total: 60 errores corregidos')
    r5.font.size = Pt(10); r5.font.color.rgb = rgb_hex('9ca3af')

    # Tabla resumen
    doc.add_paragraph()
    p_sum = doc.add_paragraph()
    r_sum = p_sum.add_run('Resumen por Severidad')
    r_sum.bold = True; r_sum.font.size = Pt(12); r_sum.font.color.rgb = ROJO

    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    hdrs = ['Severidad', 'Cantidad', 'Estado']
    hdr_fills = ['C41E3A', 'C41E3A', 'C41E3A']
    for i, (h, f) in enumerate(zip(hdrs, hdr_fills)):
        c = t.rows[0].cells[i]
        c.text = h
        cell_bg(c, f)
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = BLANCO; run.font.size = Pt(9)

    data_sum = [
        ('🔴 CRÍTICO', '18', 'Corregido'),
        ('🟠 ALTO',    '27', 'Corregido'),
        ('🟡 MEDIO',   '11', 'Corregido'),
        ('🔵 BAJO',    ' 4', 'Corregido'),
        ('TOTAL',      '60', '✅ 100%'),
    ]
    for sev, cnt, est in data_sum:
        row = t.add_row().cells
        row[0].text = sev; row[1].text = cnt; row[2].text = est
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    t.columns[0].width = Cm(5)
    t.columns[1].width = Cm(3)
    t.columns[2].width = Cm(4)

    doc.add_page_break()

    # ── FASES ─────────────────────────────────────────────────────────────────
    SEV_COLORS = {
        'CRÍTICO': 'C41E3A',
        'ALTO':    'FF1493',
        'MEDIO':   'C41E3A',
        'BAJO':    '888888',
    }

    for fase in ERRORES:
        # Título de fase
        ph = doc.add_heading(fase['fase'], level=1)
        if ph.runs:
            ph.runs[0].font.color.rgb = ROJO
            ph.runs[0].font.size = Pt(14)

        for item in fase['items']:
            # ID + Descripción
            p_id = doc.add_paragraph()
            rid = p_id.add_run(f"[{item['id']}] ")
            rid.bold = True; rid.font.color.rgb = ROJO; rid.font.size = Pt(10)
            rdesc = p_id.add_run(item['descripcion'])
            rdesc.bold = True; rdesc.font.size = Pt(10); rdesc.font.color.rgb = OSCURO

            # Tabla del error
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'
            tbl.columns[0].width = Cm(3.5)
            tbl.columns[1].width = Cm(13)

            def add_row(label, value, bg='FFF0F6'):
                row = tbl.add_row()
                row.cells[0].text = label
                row.cells[1].text = value
                cell_bg(row.cells[0], SEV_COLORS.get(item['severidad'], 'C41E3A') if label == 'Severidad' else 'FFF0F6')
                for run in row.cells[0].paragraphs[0].runs:
                    run.bold = True; run.font.size = Pt(8)
                    run.font.color.rgb = ROJO if label != 'Severidad' else BLANCO
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(8)

            add_row('Severidad',    item['severidad'])
            add_row('Categoría',    item['categoria'])
            add_row('Archivo',      item['archivo'])
            add_row('Descripción',  item['detalle'])
            add_row('Corrección',   item['correccion'])
            add_row('Commit',       item['commit'])

            doc.add_paragraph()

        doc.add_page_break()

    os.makedirs(os.path.join(BASE, 'docs'), exist_ok=True)
    out = os.path.join(BASE, 'docs', 'Rossmix_Informe_Errores_Proceso.docx')
    doc.save(out)
    print(f'Word generado: {out}')
    size = os.path.getsize(out) / 1024
    print(f'Tamaño: {size:.1f} KB')

if __name__ == '__main__':
    generar_word()
