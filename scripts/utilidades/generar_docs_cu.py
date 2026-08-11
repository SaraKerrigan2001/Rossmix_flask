"""
Genera documentación de Casos de Uso y Historias de Usuario para Rossmix.
- docs/Rossmix_CasosDeUso.docx   (Word)
- docs/Rossmix_CasosDeUso.xlsx   (Excel)
"""
import os
from datetime import date

# ── Datos de Casos de Uso ────────────────────────────────────────────────────
CASOS_DE_USO = [
    {
        "id": "CU-01", "nombre": "Registro e Inicio de Sesión",
        "actores": "Cliente, Administrador, Especialista",
        "descripcion": "Permite a los usuarios crear cuenta e iniciar sesión con rol asignado.",
        "precondicion": "El usuario no tiene cuenta activa (registro) o tiene cuenta activa (login).",
        "postcondicion": "El usuario accede al dashboard según su rol.",
        "flujo_principal": [
            "El usuario accede a /registro o /login.",
            "Ingresa nombre, email, teléfono y contraseña.",
            "El sistema valida que el email no esté registrado.",
            "El sistema crea la cuenta con tipo_usuario='cliente'.",
            "El usuario inicia sesión y es redirigido al dashboard correspondiente.",
        ],
        "flujos_alternativos": [
            "Email ya registrado → muestra error y solicita otro email.",
            "Contraseña menor a 6 caracteres → muestra validación.",
            "Cuenta desactivada → muestra mensaje de contacto al admin.",
        ],
        "reglas_negocio": [
            "La contraseña se almacena con hash Werkzeug (generate_password_hash).",
            "La sesión expira tras 8 horas de inactividad (PERMANENT_SESSION_LIFETIME).",
            "Los roles son: cliente, admin, especialista.",
        ],
    },
    {
        "id": "CU-02", "nombre": "Agendar Cita (Flujo 4 Pasos)",
        "actores": "Cliente",
        "descripcion": "El cliente agenda una cita en 4 pasos: servicio → especialista → fecha/hora → confirmación.",
        "precondicion": "El cliente tiene sesión activa.",
        "postcondicion": "La cita queda registrada con estado 'pendiente_pago' y abono de $5.000 COP.",
        "flujo_principal": [
            "Paso 1: El cliente selecciona el servicio deseado.",
            "Paso 2: Elige la especialista o selecciona asignación aleatoria.",
            "Paso 3: Selecciona fecha (máx. 90 días) y hora disponible.",
            "Paso 4: Confirma la cita; el sistema la registra con código único y token de gestión.",
            "El sistema notifica al cliente por email/WhatsApp.",
        ],
        "flujos_alternativos": [
            "Horario no disponible → muestra solo slots libres.",
            "Servicio sin empleados → mensaje de error.",
            "Fecha en el pasado → validación en frontend y backend.",
        ],
        "reglas_negocio": [
            "Abono mínimo obligatorio: $5.000 COP.",
            "Se genera código_reserva único (8 chars alfanumérico).",
            "Se genera token_gestion seguro (secrets.token_urlsafe(24)).",
            "Los slots se calculan en intervalos de 30 minutos.",
        ],
    },
    {
        "id": "CU-03", "nombre": "Cancelar Cita con Reembolso",
        "actores": "Cliente",
        "descripcion": "El cliente cancela una cita activa. Si faltan ≥2 horas, se procesa reembolso del abono.",
        "precondicion": "La cita existe con estado 'pendiente_pago' o 'confirmada'.",
        "postcondicion": "La cita queda en estado 'cancelada'. El abono se reembolsa si aplica.",
        "flujo_principal": [
            "El cliente accede a 'Mis Citas' y selecciona 'Cancelar'.",
            "El sistema calcula el tiempo restante hasta la cita.",
            "Si faltan ≥2 horas, procesa reembolso automático de $5.000 COP.",
            "La cita pasa a estado 'cancelada' y reembolsado=True.",
            "Se notifica al cliente y a los administradores.",
        ],
        "flujos_alternativos": [
            "Menos de 2 horas → cancela sin reembolso, muestra política.",
            "Cita ya cancelada → muestra error.",
        ],
        "reglas_negocio": [
            "Mínimo 2 horas de anticipación para reembolso.",
            "Monto de reembolso fijo: $5.000 COP (PasarelaPagoService).",
            "Se notifica por sistema (Notificacion) y log de WhatsApp.",
        ],
    },
    {
        "id": "CU-04", "nombre": "Reprogramar Cita",
        "actores": "Cliente, Administrador",
        "descripcion": "Reprogramación de una cita activa. El abono se transfiere a la nueva cita.",
        "precondicion": "Cita en estado 'pendiente_pago' o 'confirmada' con ≥2 horas de anticipación.",
        "postcondicion": "Cita original cancelada. Nueva cita creada con abono transferido.",
        "flujo_principal": [
            "El cliente accede desde el link de token de gestión o desde 'Mis Citas'.",
            "Selecciona nueva fecha y hora disponible.",
            "Opcionalmente cambia de especialista.",
            "El sistema cancela la cita original y crea una nueva con el abono como crédito.",
            "Se notifica al cliente con los nuevos datos.",
        ],
        "flujos_alternativos": [
            "Menos de 2 horas → reprogramación denegada.",
            "Nueva fecha no disponible → selector de otro horario.",
            "Admin reprograma por imprevisto → notifica automáticamente al cliente (CU-06).",
        ],
        "reglas_negocio": [
            "El abono previo se transfiere íntegro a la nueva cita.",
            "La cita original se marca como 'cancelada' con nota de reprogramación.",
            "Se registra en auditoria_usuarios la acción.",
        ],
    },
    {
        "id": "CU-05", "nombre": "Gestión de Agenda Diaria",
        "actores": "Administrador, Especialista",
        "descripcion": "Vista de agenda del día con cuadrícula por profesional. Permite marcar atención y liquidar citas.",
        "precondicion": "Usuario con rol admin o especialista con sesión activa.",
        "postcondicion": "Estado de citas actualizado en tiempo real.",
        "flujo_principal": [
            "Admin/especialista accede a /admin/agenda-diaria.",
            "Ve la cuadrícula de citas agrupadas por profesional para el día.",
            "Marca 'En Atención' cuando el cliente llega.",
            "Al finalizar, registra el pago del saldo pendiente (efectivo/transferencia/Nequi/Daviplata).",
            "La cita pasa a 'Completada' y se registra el pago en la tabla pagos.",
        ],
        "flujos_alternativos": [
            "Cliente no llega → marcar 'No Asistió', abono no reembolsable.",
            "Servicio adicional en curso → agregar al monto total (CU-05 ext.).",
        ],
        "reglas_negocio": [
            "Transición de estados: confirmada → en_atencion → completada.",
            "Solo se puede liquidar si está 'en_atencion'.",
            "El abono ya pagado se descuenta del saldo pendiente.",
        ],
    },
    {
        "id": "CU-06", "nombre": "Gestión de Empleados y Horarios",
        "actores": "Administrador",
        "descripcion": "CRUD de empleados, asignación de servicios, configuración de horarios semanales.",
        "precondicion": "Usuario con rol admin.",
        "postcondicion": "Empleado creado/editado con servicios y horarios actualizados.",
        "flujo_principal": [
            "Admin accede a /admin/empleados.",
            "Crea o edita empleado con nombre y especialidad.",
            "Asigna los servicios que puede realizar (M:N).",
            "Configura horarios por día de la semana (lun-vie / sáb).",
            "El empleado queda disponible para nuevas citas.",
        ],
        "flujos_alternativos": [
            "Eliminar empleado con citas futuras → desasigna citas y notifica clientes.",
            "Cambiar horario con citas afectadas → alerta de confirmación antes de aplicar.",
        ],
        "reglas_negocio": [
            "Empleado sin servicios asignados no aparece en el agendamiento.",
            "Horarios: día_semana 0=Dom, 1=Lun, ..., 6=Sáb.",
            "Al eliminar: citas futuras quedan con id_empleado=NULL para reasignación.",
        ],
    },
]

# ── Historias de Usuario ─────────────────────────────────────────────────────
HISTORIAS = [
    {
        "id": "HU-01", "cu": "CU-01", "prioridad": "Alta", "puntos": 3,
        "rol": "Cliente",
        "quiero": "registrarme con mi email y contraseña",
        "para": "acceder al sistema y agendar citas",
        "criterios": [
            "El sistema valida que el email no esté registrado.",
            "La contraseña debe tener mínimo 6 caracteres.",
            "Al registrarme exitosamente, soy redirigido al dashboard de cliente.",
            "Recibo notificación de bienvenida.",
        ],
    },
    {
        "id": "HU-02", "cu": "CU-01", "prioridad": "Alta", "puntos": 2,
        "rol": "Admin",
        "quiero": "iniciar sesión con mis credenciales",
        "para": "acceder al panel de administración",
        "criterios": [
            "El sistema verifica email y contraseña con hash.",
            "Si la cuenta está desactivada, muestra mensaje de contacto.",
            "La sesión expira tras 8 horas de inactividad.",
            "Soy redirigido al dashboard de admin.",
        ],
    },
    {
        "id": "HU-03", "cu": "CU-02", "prioridad": "Alta", "puntos": 8,
        "rol": "Cliente",
        "quiero": "agendar una cita en 4 pasos",
        "para": "reservar mi servicio de belleza en el horario que me convenga",
        "criterios": [
            "Puedo ver todos los servicios activos del salón.",
            "Puedo elegir mi especialista preferida o una aleatoria.",
            "El sistema solo muestra horarios disponibles.",
            "Recibo código de reserva y token de gestión al confirmar.",
            "El abono de $5.000 COP queda registrado.",
        ],
    },
    {
        "id": "HU-04", "cu": "CU-02", "prioridad": "Media", "puntos": 3,
        "rol": "Cliente",
        "quiero": "ver mis citas activas y pasadas",
        "para": "llevar control de mi historial en el salón",
        "criterios": [
            "Veo citas futuras con estado y botones de acción.",
            "Veo las últimas 10 citas pasadas.",
            "Puedo descargar el comprobante PDF de cada cita.",
            "Veo mi nivel de fidelidad (Bronce/Plata/Oro).",
        ],
    },
    {
        "id": "HU-05", "cu": "CU-03", "prioridad": "Alta", "puntos": 5,
        "rol": "Cliente",
        "quiero": "cancelar mi cita con anticipación",
        "para": "recuperar mi abono de $5.000 COP si cancelo a tiempo",
        "criterios": [
            "Puedo cancelar si faltan ≥2 horas para la cita.",
            "El sistema procesa el reembolso automáticamente.",
            "Recibo confirmación de cancelación y reembolso.",
            "Si cancelo con menos de 2 horas, pierdo el abono.",
        ],
    },
    {
        "id": "HU-06", "cu": "CU-04", "prioridad": "Alta", "puntos": 5,
        "rol": "Cliente",
        "quiero": "reprogramar mi cita desde el link de gestión",
        "para": "cambiar mi cita sin perder el abono pagado",
        "criterios": [
            "El link de gestión por WhatsApp/email me lleva al formulario.",
            "Puedo seleccionar nueva fecha y hora disponible.",
            "El abono previo se transfiere a la nueva cita.",
            "Recibo confirmación con los nuevos datos.",
        ],
    },
    {
        "id": "HU-07", "cu": "CU-04", "prioridad": "Alta", "puntos": 3,
        "rol": "Admin",
        "quiero": "reprogramar la cita de una cliente por imprevisto",
        "para": "gestionar cambios de agenda sin afectar la experiencia del cliente",
        "criterios": [
            "Puedo seleccionar nueva fecha y especialista.",
            "La cliente recibe notificación automática del cambio.",
            "El abono de la cliente se conserva.",
            "Queda registrado el motivo del cambio en las notas de la cita.",
        ],
    },
    {
        "id": "HU-08", "cu": "CU-05", "prioridad": "Alta", "puntos": 5,
        "rol": "Admin",
        "quiero": "ver la agenda diaria en cuadrícula por profesional",
        "para": "gestionar la operación del salón en tiempo real",
        "criterios": [
            "Veo todas las citas del día agrupadas por especialista.",
            "Puedo marcar 'En Atención' cuando el cliente llega.",
            "Puedo registrar el pago del saldo al finalizar.",
            "Veo estadísticas del día: total citas, completadas, ingresos.",
        ],
    },
    {
        "id": "HU-09", "cu": "CU-05", "prioridad": "Media", "puntos": 3,
        "rol": "Admin",
        "quiero": "marcar una cita como 'No Asistió'",
        "para": "registrar la ausencia y conservar el abono como penalidad",
        "criterios": [
            "Solo se puede marcar después de la hora de la cita.",
            "El abono de $5.000 COP no se reembolsa.",
            "La cliente puede reagendar con el abono como crédito.",
            "Recibo notificación con la opción de reagendar.",
        ],
    },
    {
        "id": "HU-10", "cu": "CU-06", "prioridad": "Alta", "puntos": 5,
        "rol": "Admin",
        "quiero": "crear y configurar empleadas con sus servicios y horarios",
        "para": "que las clientes puedan agendar citas con las especialistas disponibles",
        "criterios": [
            "Puedo crear empleada con nombre, especialidad y servicios.",
            "Configuro los horarios por día de la semana.",
            "Al eliminar, las citas futuras se desasignan y las clientes son notificadas.",
            "Puedo crear cuenta de acceso (especialista) vinculada al empleado.",
        ],
    },
    {
        "id": "HU-11", "cu": "CU-06", "prioridad": "Media", "puntos": 3,
        "rol": "Especialista",
        "quiero": "ver mis citas disponibles y aceptarlas",
        "para": "gestionar mi propia agenda de trabajo",
        "criterios": [
            "Veo las citas sin asignar para mis servicios.",
            "Puedo aceptar una cita con un clic.",
            "La cliente es notificada al asignarse la especialista.",
            "Veo mis citas próximas y mi historial del mes.",
        ],
    },
    {
        "id": "HU-12", "cu": "CU-06", "prioridad": "Baja", "puntos": 2,
        "rol": "Admin",
        "quiero": "exportar reportes a Excel por tipo y período",
        "para": "analizar el desempeño del salón fuera de la plataforma",
        "criterios": [
            "Puedo exportar citas, pagos, clientes, empleados, servicios y horarios.",
            "Los períodos disponibles son: diario, semana, mes y año.",
            "El archivo descarga automáticamente en formato .xlsx.",
            "Las columnas incluyen todos los campos relevantes con formato.",
        ],
    },
]


# ════════════════════════════════════════════════════════════════════════════
# GENERAR WORD — Casos de Uso
# ════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROJO    = RGBColor(0xC4, 0x1E, 0x3A)
BLANCO  = RGBColor(0xFF, 0xFF, 0xFF)
GRIS    = RGBColor(0xF5, 0xF5, 0xF5)
NEGRO   = RGBColor(0x1A, 0x1A, 0x1A)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p

def add_table_row(table, label, value, label_bg='C41E3A', label_fg=BLANCO):
    row = table.add_row()
    lbl = row.cells[0]
    val = row.cells[1]
    lbl.text = label
    val.text = value
    set_cell_bg(lbl, label_bg)
    for run in lbl.paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = label_fg
        run.font.size  = Pt(9)
    for run in val.paragraphs[0].runs:
        run.font.size  = Pt(9)
    lbl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    return row

def generar_word():
    doc = Document()

    # ── Márgenes ─────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Portada ───────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ROSSMIX')
    run.font.size  = Pt(32)
    run.font.bold  = True
    run.font.color.rgb = ROJO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Salón de Belleza — Sistema de Agendamiento de Citas')
    r2.font.size = Pt(14)
    r2.font.color.rgb = NEGRO

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('ESPECIFICACIÓN DE CASOS DE USO')
    r3.font.size = Pt(18)
    r3.font.bold = True

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f'Versión 1.0  ·  {date.today().strftime("%d/%m/%Y")}')
    r4.font.size  = Pt(11)
    r4.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ── Índice manual ─────────────────────────────────────────────────────────
    add_heading(doc, 'Tabla de Contenido', level=1, color=ROJO)
    for cu in CASOS_DE_USO:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f"{cu['id']} — {cu['nombre']}")
        run.font.size = Pt(11)
    doc.add_page_break()

    # ── Introducción ──────────────────────────────────────────────────────────
    add_heading(doc, '1. Introducción', level=1, color=ROJO)
    doc.add_paragraph(
        'Este documento especifica los Casos de Uso del sistema Rossmix, '
        'una aplicación web de agendamiento de citas para salón de belleza. '
        'Cubre los flujos CU-01 a CU-06 que describen las interacciones entre '
        'los actores del sistema (Cliente, Especialista, Administrador) y la aplicación Flask.'
    ).style.font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, '2. Actores del Sistema', level=2, color=ROJO)
    actores = [
        ('Cliente',       'Usuario final que agenda, cancela y reprograma citas.'),
        ('Especialista',  'Empleada del salón con acceso a su agenda y citas disponibles.'),
        ('Administrador', 'Gestiona empleados, servicios, horarios, pagos y reportes.'),
        ('Sistema',       'Procesa reembolsos, envía notificaciones y gestiona tokens.'),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = 'Actor'
    hdr[1].text = 'Descripción'
    set_cell_bg(hdr[0], 'C41E3A')
    set_cell_bg(hdr[1], 'C41E3A')
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.bold  = True
            run.font.color.rgb = BLANCO
            run.font.size  = Pt(9)
    for actor, desc in actores:
        row = t.add_row().cells
        row[0].text = actor
        row[1].text = desc
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    t.columns[0].width = Cm(4)
    t.columns[1].width = Cm(12)
    doc.add_page_break()

    # ── Casos de Uso ──────────────────────────────────────────────────────────
    add_heading(doc, '3. Casos de Uso', level=1, color=ROJO)

    for i, cu in enumerate(CASOS_DE_USO):
        add_heading(doc, f"{cu['id']} — {cu['nombre']}", level=2, color=ROJO)

        t = doc.add_table(rows=0, cols=2)
        t.style = 'Table Grid'
        t.columns[0].width = Cm(4)
        t.columns[1].width = Cm(12)

        add_table_row(t, 'ID',            cu['id'])
        add_table_row(t, 'Nombre',        cu['nombre'])
        add_table_row(t, 'Actores',       cu['actores'])
        add_table_row(t, 'Descripción',   cu['descripcion'])
        add_table_row(t, 'Precondición',  cu['precondicion'])
        add_table_row(t, 'Postcondición', cu['postcondicion'])

        # Flujo principal
        row = t.add_row()
        set_cell_bg(row.cells[0], 'C41E3A')
        p_lbl = row.cells[0].paragraphs[0]
        r_lbl = p_lbl.add_run('Flujo Principal')
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = BLANCO
        r_lbl.font.size = Pt(9)
        steps = '\n'.join(f"{j+1}. {s}" for j, s in enumerate(cu['flujo_principal']))
        row.cells[1].text = steps
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)

        # Flujos alternativos
        row2 = t.add_row()
        set_cell_bg(row2.cells[0], '7C3AED')
        r2 = row2.cells[0].paragraphs[0].add_run('Flujos\nAlternativos')
        r2.font.bold = True
        r2.font.color.rgb = BLANCO
        r2.font.size = Pt(9)
        alts = '\n'.join(f"• {a}" for a in cu['flujos_alternativos'])
        row2.cells[1].text = alts
        for run in row2.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)

        # Reglas de negocio
        row3 = t.add_row()
        set_cell_bg(row3.cells[0], '059669')
        r3 = row3.cells[0].paragraphs[0].add_run('Reglas de\nNegocio')
        r3.font.bold = True
        r3.font.color.rgb = BLANCO
        r3.font.size = Pt(9)
        rules = '\n'.join(f"• {r}" for r in cu['reglas_negocio'])
        row3.cells[1].text = rules
        for run in row3.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)

        if i < len(CASOS_DE_USO) - 1:
            doc.add_page_break()

    # ── Modelo de BD ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, '4. Modelo de Base de Datos', level=1, color=ROJO)
    doc.add_paragraph(
        'El sistema utiliza PostgreSQL con 10 tablas relacionadas mediante Foreign Keys. '
        'Los modelos ORM están implementados con SQLAlchemy en app/models/.'
    ).style.font.size = Pt(10)

    tablas_info = [
        ('usuario',             'Clientes, admins y especialistas unificados. FK→empleados.'),
        ('empleados',           'Personal del salón con especialidad y estado activo.'),
        ('servicios',           'Catálogo de servicios con precio y duración.'),
        ('empleado_servicios',  'Relación M:N entre empleados y servicios.'),
        ('horarios_empleados',  'Horarios semanales por empleado (0=Dom..6=Sáb).'),
        ('citas',               'Reservas con estados ENUM, token de gestión y código único.'),
        ('pagos',               'Transacciones asociadas a citas (1:1). ENUMs de método.'),
        ('notificaciones',      'Alertas internas para usuarios con target URL.'),
        ('auditoria_usuarios',  'Log de acciones sobre cuentas. FK→usuario (actor y afectado).'),
        ('configuraciones',     'Parámetros clave-valor del sistema. FK→usuario (creador).'),
    ]

    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    hdrs = t2.rows[0].cells
    for h, txt in zip(hdrs, ['Tabla', 'Descripción', 'FK Principal']):
        h.text = txt
        set_cell_bg(h, 'C41E3A')
        for run in h.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = BLANCO
            run.font.size = Pt(9)

    fks = {
        'usuario': '→ empleados (id_empleado)',
        'empleado_servicios': '→ empleados, → servicios',
        'horarios_empleados': '→ empleados',
        'citas': '→ usuario, → empleados, → servicios',
        'pagos': '→ citas',
        'notificaciones': '→ usuario',
        'auditoria_usuarios': '→ usuario (×2: actor y afectado)',
        'configuraciones': '→ usuario (×2: creado/modificado)',
        'empleados': '—',
        'servicios': '—',
    }
    for tabla, desc in tablas_info:
        row = t2.add_row().cells
        row[0].text = tabla
        row[1].text = desc
        row[2].text = fks.get(tabla, '—')
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    t2.columns[0].width = Cm(4)
    t2.columns[1].width = Cm(8)
    t2.columns[2].width = Cm(5)

    os.makedirs('docs', exist_ok=True)
    out = 'docs/Rossmix_CasosDeUso.docx'
    doc.save(out)
    print(f'Word guardado: {out}')


# ════════════════════════════════════════════════════════════════════════════
# GENERAR EXCEL — Casos de Uso + Historias de Usuario
# ════════════════════════════════════════════════════════════════════════════
import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter

FILL_ROJO   = PatternFill('solid', fgColor='C41E3A')
FILL_MORADO = PatternFill('solid', fgColor='7C3AED')
FILL_VERDE  = PatternFill('solid', fgColor='059669')
FILL_AZUL   = PatternFill('solid', fgColor='2563EB')
FILL_GRIS   = PatternFill('solid', fgColor='F1F5F9')
FILL_ROSA   = PatternFill('solid', fgColor='FFF0F6')
FILL_WHITE  = PatternFill('solid', fgColor='FFFFFF')

FONT_W_HDR  = Font(bold=True, color='FFFFFF', size=11, name='Calibri')
FONT_TITULO = Font(bold=True, color='C41E3A', size=16, name='Calibri')
FONT_SEC    = Font(bold=True, color='1E293B', size=10, name='Calibri')
FONT_BODY   = Font(size=10, name='Calibri')
FONT_SMALL  = Font(size=9,  name='Calibri', color='475569')

BORDER_THIN = Border(
    left=Side(style='thin', color='E2E8F0'),
    right=Side(style='thin', color='E2E8F0'),
    top=Side(style='thin', color='E2E8F0'),
    bottom=Side(style='thin', color='E2E8F0'),
)
BORDER_MED = Border(
    left=Side(style='medium', color='C41E3A'),
    right=Side(style='medium', color='C41E3A'),
    top=Side(style='medium', color='C41E3A'),
    bottom=Side(style='medium', color='C41E3A'),
)

def style_cell(cell, font=None, fill=None, align=None, border=None):
    if font:   cell.font      = font
    if fill:   cell.fill      = fill
    if align:  cell.alignment = align
    if border: cell.border    = border

def write_header_row(ws, row, cols, fill=FILL_ROJO):
    for c, (col, txt) in enumerate(cols, 1):
        cell = ws.cell(row=row, column=c, value=txt)
        style_cell(cell, font=FONT_W_HDR, fill=fill,
                   align=Alignment(horizontal='center', vertical='center', wrap_text=True),
                   border=BORDER_THIN)
        ws.column_dimensions[get_column_letter(c)].width = col

def generar_excel():
    wb = openpyxl.Workbook()

    # ── Hoja 1: Portada ───────────────────────────────────────────────────────
    ws0 = wb.active
    ws0.title = 'Portada'
    ws0.sheet_view.showGridLines = False
    ws0.column_dimensions['A'].width = 5
    ws0.column_dimensions['B'].width = 60

    ws0.row_dimensions[3].height = 40
    c = ws0.cell(row=3, column=2, value='ROSSMIX — Salón de Belleza')
    style_cell(c, font=Font(bold=True, size=24, color='C41E3A', name='Calibri'),
               align=Alignment(horizontal='left', vertical='center'))

    ws0.row_dimensions[4].height = 22
    c2 = ws0.cell(row=4, column=2, value='Casos de Uso · Historias de Usuario')
    style_cell(c2, font=Font(size=13, color='475569', name='Calibri'),
               align=Alignment(horizontal='left', vertical='center'))

    ws0.row_dimensions[5].height = 18
    c3 = ws0.cell(row=5, column=2, value=f'Versión 1.0  ·  {date.today().strftime("%d de %B de %Y")}')
    style_cell(c3, font=Font(size=10, color='94A3B8', name='Calibri'),
               align=Alignment(horizontal='left', vertical='center'))

    for r in range(8, 20):
        ws0.row_dimensions[r].height = 18
        items = [
            (8,  'Hoja',             'Contenido'),
            (9,  'Casos de Uso',     'Especificación detallada CU-01 a CU-06'),
            (10, 'Historias Usuario','Backlog con criterios de aceptación'),
            (11, 'Modelo BD',        'Tablas, columnas y relaciones FK'),
            (12, 'Estados Citas',    'Máquina de estados y transiciones'),
        ]
    for r, label, value in items:
        c_lbl = ws0.cell(row=r, column=2, value=label)
        c_val = ws0.cell(row=r, column=3, value=value)
        if r == 8:
            style_cell(c_lbl, font=FONT_W_HDR, fill=FILL_ROJO, border=BORDER_THIN,
                       align=Alignment(horizontal='center'))
            style_cell(c_val, font=FONT_W_HDR, fill=FILL_ROJO, border=BORDER_THIN,
                       align=Alignment(horizontal='center'))
        else:
            fill = FILL_ROSA if r % 2 == 0 else FILL_WHITE
            style_cell(c_lbl, font=FONT_SEC, fill=fill, border=BORDER_THIN,
                       align=Alignment(horizontal='left', vertical='center'))
            style_cell(c_val, font=FONT_BODY, fill=fill, border=BORDER_THIN,
                       align=Alignment(horizontal='left', vertical='center', wrap_text=True))
    ws0.column_dimensions['C'].width = 50

    # ── Hoja 2: Casos de Uso ──────────────────────────────────────────────────
    ws1 = wb.create_sheet('Casos de Uso')
    ws1.sheet_view.showGridLines = False

    cols_cu = [
        (8,  'ID'), (22, 'Nombre'), (14, 'Actores'),
        (35, 'Descripción'), (30, 'Precondición'), (30, 'Postcondición'),
        (45, 'Flujo Principal'), (35, 'Flujos Alternativos'), (35, 'Reglas de Negocio'),
    ]
    write_header_row(ws1, 1, cols_cu)
    ws1.row_dimensions[1].height = 28

    fills_alt = [FILL_WHITE, FILL_ROSA]
    for i, cu in enumerate(CASOS_DE_USO):
        r = i + 2
        ws1.row_dimensions[r].height = 80
        fill = fills_alt[i % 2]
        values = [
            cu['id'], cu['nombre'], cu['actores'], cu['descripcion'],
            cu['precondicion'], cu['postcondicion'],
            '\n'.join(f"{j+1}. {s}" for j, s in enumerate(cu['flujo_principal'])),
            '\n'.join(f"• {a}" for a in cu['flujos_alternativos']),
            '\n'.join(f"• {r2}" for r2 in cu['reglas_negocio']),
        ]
        for c, val in enumerate(values, 1):
            cell = ws1.cell(row=r, column=c, value=val)
            f = Font(bold=True, size=10, color='C41E3A', name='Calibri') if c == 1 else FONT_BODY
            style_cell(cell, font=f, fill=fill, border=BORDER_THIN,
                       align=Alignment(horizontal='left', vertical='top', wrap_text=True))

    ws1.freeze_panes = 'A2'

    # ── Hoja 3: Historias de Usuario ──────────────────────────────────────────
    ws2 = wb.create_sheet('Historias Usuario')
    ws2.sheet_view.showGridLines = False

    cols_hu = [
        (8,  'ID'), (8,  'CU'), (12, 'Prioridad'), (8, 'Puntos'),
        (14, 'Rol'), (30, 'Como...'), (35, 'Quiero...'),
        (35, 'Para...'), (50, 'Criterios de Aceptación'),
    ]
    write_header_row(ws2, 1, cols_hu, fill=FILL_MORADO)
    ws2.row_dimensions[1].height = 28

    prio_colors = {'Alta': 'FEE2E2', 'Media': 'FEF3C7', 'Baja': 'DCFCE7'}
    for i, hu in enumerate(HISTORIAS):
        r = i + 2
        ws2.row_dimensions[r].height = 70
        prio_fill = PatternFill('solid', fgColor=prio_colors.get(hu['prioridad'], 'FFFFFF'))
        base_fill = FILL_ROSA if i % 2 == 0 else FILL_WHITE
        values = [
            hu['id'], hu['cu'], hu['prioridad'], hu['puntos'],
            hu['rol'],
            f"Como {hu['rol']}",
            hu['quiero'],
            hu['para'],
            '\n'.join(f"✓ {c}" for c in hu['criterios']),
        ]
        for c, val in enumerate(values, 1):
            cell = ws2.cell(row=r, column=c, value=val)
            if c == 1:
                f = Font(bold=True, size=10, color='7C3AED', name='Calibri')
                fill = base_fill
            elif c == 3:
                f = Font(bold=True, size=10, name='Calibri')
                fill = prio_fill
            elif c == 4:
                f = Font(bold=True, size=11, color='2563EB', name='Calibri')
                fill = base_fill
            else:
                f = FONT_BODY
                fill = base_fill
            style_cell(cell, font=f, fill=fill, border=BORDER_THIN,
                       align=Alignment(horizontal='left', vertical='top', wrap_text=True))

    ws2.freeze_panes = 'A2'

    # ── Hoja 4: Modelo BD ─────────────────────────────────────────────────────
    ws3 = wb.create_sheet('Modelo BD')
    ws3.sheet_view.showGridLines = False

    cols_bd = [
        (20, 'Tabla'), (18, 'Columna'), (20, 'Tipo'), (12, 'PK/FK/UK'),
        (10, 'Nullable'), (20, 'Default'), (40, 'Descripción'),
    ]
    write_header_row(ws3, 1, cols_bd, fill=FILL_VERDE)
    ws3.row_dimensions[1].height = 28

    BD_COLS = [
        # tabla, columna, tipo, clave, nullable, default, descripcion
        ('usuario','id','SERIAL','PK','NO','—','Clave primaria autoincremental'),
        ('usuario','nombre','VARCHAR(100)','','NO','—','Nombre completo'),
        ('usuario','email','VARCHAR(150)','UK','NO','—','Email único del usuario'),
        ('usuario','telefono','VARCHAR(20)','','NO','—','Celular/WhatsApp'),
        ('usuario','password','VARCHAR(200)','','NO','—','Hash Werkzeug'),
        ('usuario','tipo_usuario','VARCHAR(20)','','NO','cliente','admin|cliente|especialista'),
        ('usuario','activo','BOOLEAN','','NO','TRUE','Cuenta activa/desactivada'),
        ('usuario','id_empleado','INTEGER','FK→empleados','SÍ','NULL','Vínculo para especialistas'),
        ('empleados','id_empleado','SERIAL','PK','NO','—','Clave primaria'),
        ('empleados','nombre','VARCHAR(100)','','NO','—','Nombre de la especialista'),
        ('empleados','especialidad','VARCHAR(100)','','SÍ','NULL','Área de especialización'),
        ('empleados','activo','BOOLEAN','','NO','TRUE','Disponible para citas'),
        ('servicios','id_servicio','SERIAL','PK','NO','—','Clave primaria'),
        ('servicios','nombre_servicio','VARCHAR(100)','','NO','—','Nombre del servicio'),
        ('servicios','precio_total','NUMERIC(10,2)','','NO','—','Precio en COP'),
        ('servicios','duracion_minutos','INTEGER','','NO','—','Duración del servicio'),
        ('empleado_servicios','id_empleado','INTEGER','FK→empleados','NO','—','Parte de PK compuesta'),
        ('empleado_servicios','id_servicio','INTEGER','FK→servicios','NO','—','Parte de PK compuesta'),
        ('horarios_empleados','id_horario','SERIAL','PK','NO','—','Clave primaria'),
        ('horarios_empleados','id_empleado','INTEGER','FK→empleados','NO','—','Empleado asociado'),
        ('horarios_empleados','dia_semana','INTEGER','','NO','—','0=Dom 1=Lun...6=Sáb'),
        ('horarios_empleados','hora_inicio','TIME','','NO','—','Hora de inicio del turno'),
        ('horarios_empleados','hora_fin','TIME','','NO','—','Hora de fin del turno'),
        ('citas','id_cita','SERIAL','PK','NO','—','Clave primaria'),
        ('citas','id_cliente','INTEGER','FK→usuario','NO','—','CASCADE al borrar usuario'),
        ('citas','id_empleado','INTEGER','FK→empleados','SÍ','NULL','SET NULL al borrar empleado'),
        ('citas','id_servicio','INTEGER','FK→servicios','NO','—','RESTRICT al borrar servicio'),
        ('citas','estado','estado_cita_enum','','NO','pendiente_pago','Ver máquina de estados'),
        ('citas','monto_abono','NUMERIC(10,2)','','SÍ','5000','Abono mínimo $5.000 COP'),
        ('citas','codigo_reserva','VARCHAR(20)','UK','SÍ','NULL','Generado por CitaService'),
        ('citas','token_gestion','VARCHAR(32)','UK','SÍ','NULL','secrets.token_urlsafe(24)'),
        ('pagos','id_pago','SERIAL','PK','NO','—','Clave primaria'),
        ('pagos','id_cita','INTEGER','FK→citas UK','NO','—','1 pago por cita (CASCADE)'),
        ('pagos','monto','NUMERIC(10,2)','','NO','—','Monto pagado'),
        ('pagos','metodo_pago','metodo_pago_enum','','NO','efectivo','Ver ENUM'),
        ('notificaciones','id','SERIAL','PK','NO','—','Clave primaria'),
        ('notificaciones','id_usuario','INTEGER','FK→usuario','NO','—','CASCADE al borrar usuario'),
        ('notificaciones','titulo','VARCHAR(200)','','NO','—','Asunto de la notificación'),
        ('notificaciones','leido','BOOLEAN','','NO','FALSE','Estado de lectura'),
        ('auditoria_usuarios','id','SERIAL','PK','NO','—','Clave primaria'),
        ('auditoria_usuarios','id_usuario','INTEGER','FK→usuario','SÍ','NULL','SET NULL al borrar'),
        ('auditoria_usuarios','id_actor','INTEGER','FK→usuario','SÍ','NULL','Admin que actuó'),
        ('auditoria_usuarios','accion','VARCHAR(50)','','NO','—','login|editar|eliminar|etc.'),
        ('auditoria_usuarios','ip_address','VARCHAR(45)','','SÍ','NULL','IPv4 o IPv6'),
        ('configuraciones','id','SERIAL','PK','NO','—','Clave primaria'),
        ('configuraciones','clave','VARCHAR(120)','UK','NO','—','Identificador del parámetro'),
        ('configuraciones','valor','TEXT','','NO','—','Valor en texto plano'),
        ('configuraciones','creado_por','INTEGER','FK→usuario','SÍ','NULL','SET NULL al borrar'),
        ('configuraciones','modificado_por','INTEGER','FK→usuario','SÍ','NULL','SET NULL al borrar'),
    ]

    tabla_fills = {
        'usuario': 'FFF0F6', 'empleados': 'F0FDF4', 'servicios': 'F0FDF4',
        'empleado_servicios': 'EFF6FF', 'horarios_empleados': 'EFF6FF',
        'citas': 'FFF0F6', 'pagos': 'FFFBEB',
        'notificaciones': 'FDF4FF', 'auditoria_usuarios': 'F8FAFC',
        'configuraciones': 'F8FAFC',
    }
    for i, row_data in enumerate(BD_COLS):
        r = i + 2
        ws3.row_dimensions[r].height = 18
        tabla = row_data[0]
        fill = PatternFill('solid', fgColor=tabla_fills.get(tabla, 'FFFFFF'))
        for c, val in enumerate(row_data, 1):
            cell = ws3.cell(row=r, column=c, value=val)
            if c == 1:
                f = Font(bold=True, size=9, color='059669', name='Calibri')
            elif c == 4 and val:
                f = Font(bold=True, size=9, color='C41E3A' if 'PK' in val else '2563EB', name='Calibri')
            else:
                f = Font(size=9, name='Calibri')
            style_cell(cell, font=f, fill=fill, border=BORDER_THIN,
                       align=Alignment(horizontal='left', vertical='center', wrap_text=True))
    ws3.freeze_panes = 'A2'

    # ── Hoja 5: Estados de Citas ──────────────────────────────────────────────
    ws4 = wb.create_sheet('Estados Citas')
    ws4.sheet_view.showGridLines = False

    cols_est = [(20,'Estado'), (20,'Desde'), (20,'Hacia'), (15,'Quién'), (40,'Descripción')]
    write_header_row(ws4, 1, cols_est, fill=FILL_AZUL)

    transiciones = [
        ('pendiente_pago',  '—',             'confirmada',    'Admin',       'Admin confirma el abono del cliente'),
        ('pendiente_pago',  '—',             'cancelada',     'Cliente',     'Cliente cancela antes de 2h → reembolso'),
        ('confirmada',      'pendiente_pago','en_atencion',   'Admin/Espec', 'Cliente llega al salón'),
        ('confirmada',      'pendiente_pago','cancelada',     'Cliente',     'Cancelación con ≥2h de anticipación'),
        ('confirmada',      'pendiente_pago','no_asistio',    'Admin',       'Hora pasada, cliente no llegó'),
        ('en_atencion',     'confirmada',    'completada',    'Admin/Espec', 'Servicio finalizado y saldo cobrado'),
        ('completada',      'en_atencion',   '—',             '—',           'Estado terminal. No hay transición.'),
        ('cancelada',       'cualquiera',    '—',             '—',           'Estado terminal (abono reembolsado o no)'),
        ('no_asistio',      'confirmada',    '(reagendar)',   'Cliente',     'Cliente puede reagendar con abono como crédito'),
    ]
    estado_colors = {
        'pendiente_pago': 'FEF3C7', 'confirmada': 'DCFCE7',
        'en_atencion': 'DBEAFE', 'completada': 'D1FAE5',
        'cancelada': 'FEE2E2', 'no_asistio': 'F3F4F6',
    }
    for i, (est, desde, hacia, quien, desc) in enumerate(transiciones):
        r = i + 2
        ws4.row_dimensions[r].height = 22
        fill = PatternFill('solid', fgColor=estado_colors.get(est, 'FFFFFF'))
        for c, val in enumerate([est, desde, hacia, quien, desc], 1):
            cell = ws4.cell(row=r, column=c, value=val)
            f = Font(bold=True, size=10, name='Calibri') if c == 1 else FONT_BODY
            style_cell(cell, font=f, fill=fill, border=BORDER_THIN,
                       align=Alignment(horizontal='left', vertical='center', wrap_text=True))
    ws4.freeze_panes = 'A2'

    # ── Guardar ───────────────────────────────────────────────────────────────
    os.makedirs('docs', exist_ok=True)
    out = 'docs/Rossmix_CasosDeUso.xlsx'
    wb.save(out)
    print(f'Excel guardado: {out}')


# ── Ejecutar ──────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    generar_word()
    generar_excel()
    print('\nDocumentos generados en docs/')
