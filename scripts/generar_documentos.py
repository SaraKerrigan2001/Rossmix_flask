"""
Genera Excel y Word actualizados con TODAS las HU y CU del proyecto Rossmix.
Incluye las 6 originales + 9 nuevas (especialistas, admin, auditoría, etc.)
"""
import os
from openpyxl import Workbook
from openpyxl.styles import (PatternFill, Font, Alignment, Border, Side,
                               GradientFill)
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTDIR = os.path.join(BASE, 'Rossmix-20260810T155230Z-1-001', 'Rossmix',
                      'Caso de Uso - Historia de Usuario')
os.makedirs(OUTDIR, exist_ok=True)

# ── Paleta de colores ─────────────────────────────────────────────────────────
ROSA       = "C41E3A"
ROSA_LIGHT = "FFD6E8"
ROSA_MED   = "FF6B9D"
VERDE      = "059669"
VERDE_L    = "D1FAE5"
AZUL       = "2563EB"
AZUL_L     = "DBEAFE"
MORADO     = "7C3AED"
MORADO_L   = "EDE9FF"
NARANJA    = "D97706"
NARANJA_L  = "FEF3C7"
GRIS       = "475569"
GRIS_L     = "F1F5F9"
BLANCO     = "FFFFFF"
NEGRO      = "1A1A2E"

# ── Datos completos: HU + CU ─────────────────────────────────────────────────
# Formato: (id, titulo, actor, prioridad, pts, como, quiero, para_que,
#           criterios[], notas)
HISTORIAS = [
    # ── ORIGINALES (actualizadas) ─────────────────────────────────────────────
    ("HU-01","Consultar Servicios y Disponibilidad","Cliente","Alta",5,
     "clienta del salón Rossmix",
     "ver el catálogo de servicios con precios, duración e imágenes reales",
     "elegir el servicio que mejor se adapte a mis necesidades antes de agendar",
     ["La página de inicio muestra galería de servicios con imagen, precio y duración",
      "Se puede filtrar por categoría (uñas, cabello, depilación, cejas/pestañas)",
      "Cada servicio muestra las especialistas que lo realizan",
      "El botón 'Agendar' lleva directo al paso 1 del flujo de citas",
      "El catálogo es visible sin necesidad de iniciar sesión"],
     "Implementado en app/templates/index.html con galería de servicios"),

    ("HU-02","Agendar Cita y Pagar Abono","Cliente","Alta",8,
     "clienta registrada en Rossmix",
     "agendar una cita en 4 pasos: servicio → especialista → fecha/hora → confirmación",
     "reservar mi turno con abono de $5.000 COP que garantiza mi lugar",
     ["Paso 1: seleccionar servicio activo del catálogo",
      "Paso 2: elegir especialista disponible o 'Asignación automática'",
      "Paso 3: calendario con horarios disponibles según agenda de la especialista",
      "Paso 4: resumen completo antes de confirmar con código de reserva único",
      "El sistema genera código_reserva y token_gestion al confirmar",
      "El abono de $5.000 COP se registra en tabla pagos",
      "La cita queda en estado 'pendiente_pago' hasta confirmar el pago",
      "La clienta recibe notificación interna de confirmación"],
     "Flujo implementado en app/views/citas.py pasos 1-4 + confirmar_cita"),

    ("HU-03","Cancelar Cita con Devolución","Cliente","Media",5,
     "clienta con cita confirmada",
     "cancelar mi cita desde 'Mis Citas' con política clara de reembolso",
     "gestionar mi agenda de manera flexible sin perder mi dinero si cancelo a tiempo",
     ["Solo se puede cancelar con ≥2 horas de anticipación",
      "Con ≥2h: cita pasa a 'cancelada' y se procesa reembolso del abono",
      "Con <2h: el sistema informa que el abono no es reembolsable",
      "Si no asistió: badge 'no_asistio' + mensaje de política de abono no reembolsable",
      "Opción de reagendar desde el historial de citas con crédito del abono",
      "El admin recibe notificación automática de la cancelación"],
     "Implementado en app/views/citas.py :: cancelar_cita + reagendar_no_asistio"),

    ("HU-04","Gestionar Horarios y Personal","Administrador","Alta",5,
     "administradora del salón",
     "crear, editar y eliminar empleadas con sus horarios y servicios asignados",
     "mantener actualizada la disponibilidad del equipo para que los agendamientos sean correctos",
     ["CRUD completo de empleadas con nombre y especialidad",
      "Asignación de servicios por empleada (many-to-many)",
      "Configuración de horarios por día de la semana con hora inicio/fin",
      "Al eliminar empleada: lista de clientas afectadas + opción de reasignar",
      "Las citas de la empleada eliminada quedan con id_empleado=NULL para reasignar",
      "Modales elegantes para crear/editar sin salir de la página"],
     "Implementado en app/views/admin/empleados.py + horarios.py"),

    ("HU-05","Gestionar Agenda del Día","Administrador","Alta",8,
     "administradora del salón",
     "ver todas las citas del día actual con estado, especialista y pago",
     "coordinar el trabajo del salón en tiempo real y registrar pagos pendientes",
     ["Vista de agenda diaria filtrando solo citas de hoy",
      "Columnas: hora, clienta, servicio, especialista, estado, monto, acciones",
      "Cambiar estado de cita directamente con select (sin recargar página)",
      "Botón 'Registrar Pago' para citas sin pago completado",
      "Botón 'Asignar' para citas sin especialista asignada",
      "Exportar a Excel por período (diario, semanal, mensual, anual)"],
     "Implementado en app/views/admin/dashboard.py :: agenda_diaria"),

    ("HU-06","Reprogramar / Modificar Cita","Cliente/Admin","Media",5,
     "clienta o administradora",
     "reprogramar una cita existente a una nueva fecha y hora disponible",
     "flexibilizar mi agenda sin perder el abono ya pagado",
     ["Solo se puede reprogramar con ≥2 horas de anticipación",
      "El sistema reutiliza el paso 3 (fecha/hora) del flujo de agendamiento",
      "El token_gestion de la cita original permite acceder sin iniciar sesión",
      "El admin puede reprogramar desde Gestión de Citas con botón Reasignar",
      "La clienta recibe notificación de la reprogramación",
      "El abono se mantiene vinculado a la nueva cita"],
     "Implementado en app/views/citas.py :: reprogramar_cita_form + gestionar_cita"),

    # ── NUEVAS ────────────────────────────────────────────────────────────────
    ("HU-07","Portal de Especialista — Ver y Aceptar Citas","Especialista","Alta",8,
     "especialista del salón con cuenta de acceso",
     "ver en tiempo real las citas disponibles (sin especialista asignada) y aceptarlas",
     "gestionar mi agenda de forma autónoma y tomar las citas que me convengan",
     ["La especialista inicia sesión con email/contraseña (tipo_usuario='especialista')",
      "Dashboard propio con: mis citas próximas, citas disponibles, completadas del mes",
      "Las citas disponibles se muestran como tarjetas con servicio, clienta, fecha y hora",
      "Botón '¡Aceptar!' asigna la cita a la especialista en tiempo real",
      "Si otra especialista ya tomó la cita: mensaje 'ya fue tomada'",
      "Solo se muestran citas para servicios que la especialista realiza",
      "La clienta recibe notificación automática '¡Tu especialista fue asignada!'"],
     "Implementado en app/views/especialista.py :: dashboard + citas_disponibles + aceptar_cita"),

    ("HU-08","Gestión de Cuentas de Especialistas","Administrador","Alta",5,
     "administradora del salón",
     "crear y gestionar cuentas de acceso al portal para cada especialista",
     "controlar quién puede acceder al sistema y vincular cada cuenta con su perfil de empleada",
     ["Sección 'Especialistas' en el panel admin lista todas las cuentas activas",
      "Botón 'Crear Acceso' abre modal con: seleccionar empleada, email, contraseña",
      "Se muestra qué empleadas aún no tienen cuenta de acceso",
      "El admin puede resetear la contraseña de cualquier especialista desde modal",
      "Eliminar cuenta no elimina a la empleada, solo su acceso al portal",
      "Solo empleadas activas pueden tener cuenta de especialista"],
     "Implementado en app/views/admin/especialistas.py"),

    ("HU-09","Distribución de Citas por el Admin","Administrador","Alta",8,
     "administradora del salón",
     "asignar citas sin especialista a las profesionales disponibles desde un panel centralizado",
     "garantizar que ninguna cita quede sin atender y distribuir la carga equitativamente",
     ["Panel 'Distribución de Citas' muestra todas las citas sin especialista asignada",
      "Cada fila tiene un selector de especialista con su carga de trabajo actual",
      "Botón 'Auto-distribuir' asigna automáticamente según quien tenga menos citas",
      "Botón individual por cita para asignación una a una",
      "Botón 'Guardar Todas' procesa las asignaciones en batch con una sola petición",
      "Barra de carga visual por especialista (verde/amarillo/rojo)",
      "La clienta recibe notificación al asignarle una especialista"],
     "Implementado en app/views/admin/citas.py :: citas_distribucion + citas_asignar_batch"),

    ("HU-10","Reasignar Especialista a una Cita","Administrador","Media",5,
     "administradora del salón",
     "cambiar la especialista asignada a una cita existente desde Gestión de Citas",
     "corregir asignaciones incorrectas o cubrir ausencias del personal",
     ["Botón 'Reasignar' visible en citas con y sin especialista asignada",
      "Modal muestra lista de especialistas que realizan el servicio de la cita",
      "Al confirmar: la cita se actualiza y pasa a estado 'confirmada'",
      "La clienta recibe notificación de la nueva especialista asignada",
      "La acción queda registrada en la auditoría del sistema"],
     "Implementado en app/views/admin/citas.py :: citas_reasignar_empleado"),

    ("HU-11","Gestión de Pagos y Reembolsos","Administrador","Alta",5,
     "administradora del salón",
     "registrar pagos completos, ver pagos por confirmar y procesar reembolsos",
     "mantener el control financiero del salón con registro detallado de cada transacción",
     ["Vista de todos los pagos con: cliente, servicio, monto, método, estado, fecha",
      "Sección 'Pagos por Confirmar' lista citas activas sin pago registrado",
      "Formulario de registro de pago: monto, método (efectivo/tarjeta/transferencia/Nequi/Daviplata)",
      "Modal de reembolso con monto destacado y confirmación obligatoria",
      "Al reembolsar: pago se elimina y cita pasa a 'cancelada'",
      "Exportación a Excel por período"],
     "Implementado en app/views/admin/pagos.py"),

    ("HU-12","Auditoría en Tiempo Real de Usuarios","Administrador","Media",3,
     "administradora del salón",
     "ver un registro automático de cada nuevo usuario registrado en el sistema",
     "tener trazabilidad completa del crecimiento de la base de clientes",
     ["Cada nuevo registro en 'usuario' se captura automáticamente vía trigger PostgreSQL",
      "Tabla auditoria_usuarios guarda: nombre, email, tipo, fecha, acción",
      "El trigger emite pg_notify al canal 'nuevo_usuario' para integración en tiempo real",
      "La vista vista_nuevos_usuarios permite consultar desde pgAdmin",
      "Backfill disponible para usuarios creados antes del trigger",
      "Sin impacto en rendimiento — el trigger es AFTER INSERT asíncrono"],
     "Implementado en Rossmix.sql sección 19 + app/views/admin/especialistas.py"),

    ("HU-13","Notificaciones Internas","Cliente/Admin","Media",5,
     "usuaria del sistema (cliente o admin)",
     "recibir notificaciones internas dentro de la aplicación sobre mis citas y eventos importantes",
     "estar siempre informada sin necesidad de revisar mi correo electrónico",
     ["Badge contador en el ícono de campana del navbar actualizado en tiempo real",
      "Lista de notificaciones con: título, mensaje, fecha y estado leído/no leído",
      "Click en notificación marca como leída y redirige al target correspondiente",
      "Botón 'Marcar todas como leídas'",
      "Notificaciones automáticas en: confirmación de cita, asignación de especialista, cancelación, reembolso",
      "Admin recibe notificaciones de: cancelaciones de clientas, nuevos registros"],
     "Implementado en app/views/notificaciones.py + app/utils/helpers.py :: add_notificacion"),

    ("HU-14","Descarga de Comprobante PDF","Cliente","Baja",3,
     "clienta con cita confirmada",
     "descargar un comprobante PDF con los detalles completos de mi cita",
     "tener un documento oficial de mi reserva para mostrar en el salón o guardar como respaldo",
     ["El PDF incluye: código de reserva, servicio, especialista, fecha/hora, montos",
      "Diseño con colores Rossmix (rosa #C41E3A) y tabla formateada",
      "Disponible desde 'Mis Citas' y desde la página de confirmación",
      "Nombre de archivo: rossmix_cita_XXXXXX.pdf",
      "Se genera con ReportLab en tiempo real sin almacenar en servidor"],
     "Implementado en app/views/citas.py :: descargar_cita_pdf"),

    ("HU-15","Gestión de Clientes","Administrador","Media",5,
     "administradora del salón",
     "ver, editar y gestionar el perfil de todas las clientas registradas",
     "mantener la base de datos de clientas actualizada y acceder a su historial",
     ["Lista de clientes con: nombre, email, teléfono, total citas, cancelaciones, estado",
      "Filtro para ver clientes de 'Hoy' vs 'Todos'",
      "Modal de edición: nombre, email, teléfono, cambio de contraseña, estado activo",
      "Botón historial enlaza a las citas de esa clienta filtradas",
      "Export a Excel por período",
      "No se puede eliminar si tiene citas futuras confirmadas"],
     "Implementado en app/views/admin/clientes.py"),
]

# ── Casos de Uso (técnico) ────────────────────────────────────────────────────
CASOS_USO = [
    ("CU-01","Consultar Servicios y Disponibilidad","Cliente","app/templates/index.html",
     "El cliente accede a la URL raíz del sistema",
     ["Actor accede a / (página de inicio)",
      "Sistema muestra galería de servicios con imagen, nombre, precio y duración",
      "Actor puede ver servicios sin autenticación",
      "Actor hace click en 'Agendar' y es redirigido al flujo de citas"],
     ["No hay servicios activos → mensaje 'Próximamente'",
      "Actor no registrado → redirigir a /registro al intentar agendar"]),

    ("CU-02","Agendar Cita y Pagar Abono","Cliente","app/views/citas.py",
     "Cliente autenticado con tipo_usuario='cliente'",
     ["GET /citas/agendar/paso1 → lista servicios activos",
      "GET /citas/agendar/paso2/<id_servicio> → lista empleadas del servicio",
      "GET /citas/agendar/paso3/<id_servicio>/<id_empleado> → calendario disponible",
      "GET /citas/horarios-disponibles?fecha&id_empleado&id_servicio → JSON slots",
      "GET /citas/agendar/paso4 → resumen de confirmación",
      "POST /citas/confirmar → crea Cita + Pago, genera codigo_reserva y token_gestion",
      "Sistema emite notificación interna a la clienta",
      "Redirección a /citas/confirmada/<codigo>"],
     ["id_empleado=0 → asignación aleatoria entre empleadas del servicio",
      "Sin slots disponibles → mensaje 'Sin horarios disponibles para esta fecha'",
      "Fecha en el pasado → error 'No puedes agendar en el pasado'"]),

    ("CU-03","Cancelar Cita","Cliente","app/views/citas.py :: cancelar_cita",
     "Cliente autenticado, cita en estado pendiente_pago o confirmada",
     ["Cliente accede a /citas/mis-citas",
      "Sistema muestra citas futuras con tiempo restante",
      "POST /citas/cancelar/<id_cita> con tiempo_restante ≥ 2h",
      "Sistema actualiza estado='cancelada', procesa reembolso",
      "Sistema notifica a clienta y admins"],
     ["tiempo_restante < 2h → HTTP 400 'Debes cancelar con mínimo 2h de anticipación'",
      "Cita no pertenece al usuario → HTTP 404",
      "No asistió → estado='no_asistio', abono no reembolsable, opción de reagendar"]),

    ("CU-04","Gestionar Empleadas y Horarios","Administrador","app/views/admin/empleados.py + horarios.py",
     "Usuario autenticado con tipo_usuario='admin'",
     ["GET /admin/empleados → lista con modal crear/editar",
      "POST /admin/empleados/crear → crea empleada con servicios asignados (AJAX)",
      "GET /admin/empleados/datos/<id> → JSON con datos para modal editar",
      "POST /admin/empleados/editar/<id> → actualiza empleada (AJAX)",
      "GET /admin/empleados/clientes-afectados/<id> → JSON citas futuras",
      "POST /admin/empleados/eliminar/<id> → verifica afectadas + elimina",
      "GET /admin/horarios → vista por empleada con barras de horario",
      "POST /admin/horarios/crear/<id_empleado> → nuevo slot horario (AJAX)"],
     ["Empleada tiene citas futuras → mostrar lista afectadas antes de eliminar",
      "Horario duplicado para mismo día → HTTP 400",
      "hora_inicio >= hora_fin → validación rechazada"]),

    ("CU-05","Gestión de Agenda Diaria","Administrador","app/views/admin/dashboard.py :: agenda_diaria",
     "Admin autenticado",
     ["GET /admin/agenda-diaria → lista citas del día actual (LEFT JOIN empleados)",
      "POST /admin/citas/cambiar-estado/<id> → actualiza estado via AJAX",
      "GET /admin/pagos/confirmar → citas activas sin pago",
      "GET /admin/citas/distribucion → panel de asignación de citas sin especialista",
      "POST /admin/citas/asignar-batch → asigna múltiples citas en batch JSON"],
     ["Sin citas hoy → mensaje 'No hay citas programadas para hoy'",
      "Estado inválido → HTTP 400"]),

    ("CU-06","Reprogramar / Modificar Cita","Cliente/Admin","app/views/citas.py :: reprogramar_cita_form",
     "Clienta autenticada o Admin; cita con tiempo_restante ≥ 2h",
     ["GET /citas/reprogramar/<id_cita> → reutiliza template paso3_fecha_hora",
      "Clienta selecciona nueva fecha y hora disponible",
      "Sistema valida disponibilidad del empleado",
      "POST /citas/confirmar con id_cita_original → actualiza fecha_hora_inicio/fin",
      "Sistema notifica a clienta de la reprogramación"],
     ["Sin slots disponibles → mensaje y redirigir",
      "tiempo_restante < 2h → rechazar"]),

    ("CU-07","Aceptar Cita Disponible (Especialista)","Especialista","app/views/especialista.py",
     "Usuario con tipo_usuario='especialista' vinculado a empleada activa",
     ["GET /especialista/dashboard → citas disponibles + mis citas próximas",
      "GET /especialista/citas-disponibles → listado completo por servicio",
      "POST /especialista/aceptar-cita/<id_cita> → asigna id_empleado atomicamente",
      "Sistema verifica que id_empleado IS NULL antes de asignar",
      "Sistema notifica a clienta de la asignación",
      "Cita pasa de pendiente_pago a confirmada si aplica"],
     ["Cita ya tomada → HTTP 409 'Esta cita ya fue tomada por otra especialista'",
      "Servicio no es de la especialista → HTTP 403",
      "Especialista sin vínculo a empleada → redirigir a login"]),

    ("CU-08","Gestionar Cuentas Especialistas","Administrador","app/views/admin/especialistas.py",
     "Admin autenticado",
     ["GET /admin/especialistas → lista cuentas + empleadas sin cuenta",
      "POST /admin/especialistas/crear → crea usuario tipo_usuario='especialista' con id_empleado",
      "POST /admin/especialistas/reset-password/<id> → actualiza password hash",
      "POST /admin/especialistas/eliminar/<id> → elimina usuario (no la empleada)"],
     ["Email duplicado → HTTP 400 'El email ya está registrado'",
      "Contraseña < 6 chars → validación rechazada",
      "id_empleado no es de empleada activa → error"]),

    ("CU-09","Distribución Masiva de Citas","Administrador","app/views/admin/citas.py",
     "Admin autenticado; existen citas con id_empleado IS NULL",
     ["GET /admin/citas/distribucion → citas sin asignar + carga por especialista",
      "Admin selecciona especialista por cada cita o usa 'Auto-distribuir'",
      "Auto-distribuir → asigna según menor carga (menos citas futuras activas)",
      "POST /admin/citas/asignar-batch con JSON [{id_cita, id_empleado}]",
      "Sistema actualiza todas las asignaciones + notifica a cada clienta"],
     ["Sin citas pendientes → panel vacío con mensaje '¡Todo asignado!'",
      "id_empleado inválido → se registra en lista de errores de respuesta"]),

    ("CU-10","Descargar Comprobante PDF","Cliente","app/views/citas.py :: descargar_cita_pdf",
     "Cliente autenticado con cita propia de cualquier estado",
     ["GET /citas/descargar-pdf/<id_cita>",
      "Sistema verifica que la cita pertenece al usuario (o es admin)",
      "ReportLab genera PDF en memoria con tabla formateada",
      "Response con Content-Type application/pdf + nombre de archivo"],
     ["ReportLab no instalado → flash error instructivo",
      "Cita de otro usuario → HTTP 403 redirect a mis_citas",
      "Error de generación → flash con detalle del error"]),

    ("CU-11","Registrar y Gestionar Pagos","Administrador","app/views/admin/pagos.py",
     "Admin autenticado",
     ["GET /admin/pagos → lista todos los pagos con filtros",
      "GET /admin/pagos/confirmar → citas sin pago completado",
      "GET /admin/pagos/registrar/<id_cita> → formulario de pago",
      "POST /admin/pagos/registrar/<id_cita> → crea Pago + actualiza saldo_pendiente",
      "POST /admin/pagos/eliminar/<id_pago> → reembolso: elimina pago + cancela cita"],
     ["Cita ya tiene pago → redirigir a lista de pagos",
      "Monto inválido → validación rechazada",
      "Cita cancelada → no se puede registrar pago"]),

    ("CU-12","Recibir Notificaciones Internas","Cliente/Admin","app/views/notificaciones.py",
     "Usuario autenticado de cualquier tipo",
     ["Context processor inject_notificaciones carga contador en cada request",
      "GET /notificaciones → lista todas las notificaciones del usuario",
      "POST /notificaciones/marcar-leida/<id> → marca una como leída",
      "POST /notificaciones/marcar-todas → marca todas como leídas",
      "add_notificacion() se llama en: confirmar_cita, cancelar_cita, aceptar_cita, reasignar"],
     ["Sin notificaciones → mensaje 'No tienes notificaciones'",
      "Notificación de otro usuario → HTTP 403"]),
]


# ── Helpers de estilo Excel ───────────────────────────────────────────────────
def fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def font(bold=False, size=10, color="000000", italic=False):
    return Font(name="Calibri", bold=bold, size=size,
                color=color, italic=italic)

def border_thin():
    side = Side(style='thin', color="D1D5DB")
    return Border(left=side, right=side, top=side, bottom=side)

def align(h='left', v='center', wrap=True):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

def set_col_width(ws, col_letter, width):
    ws.column_dimensions[col_letter].width = width

def merge_title(ws, cell_range, text, bg, fg="FFFFFF", size=13):
    ws.merge_cells(cell_range)
    c = ws[cell_range.split(':')[0]]
    c.value = text
    c.fill  = fill(bg)
    c.font  = font(True, size, fg)
    c.alignment = align('center', 'center')

def header_row(ws, row_num, headers, widths, bg=ROSA, fg=BLANCO):
    for i, (h, w) in enumerate(zip(headers, widths), 1):
        c = ws.cell(row=row_num, column=i, value=h)
        c.fill  = fill(bg)
        c.font  = font(True, 10, fg)
        c.alignment = align('center', 'center')
        c.border = border_thin()
        ws.column_dimensions[get_column_letter(i)].width = w

def data_cell(ws, row, col, val, bg=BLANCO, bold=False, center=False):
    c = ws.cell(row=row, column=col, value=val)
    c.fill      = fill(bg)
    c.font      = font(bold, 10)
    c.alignment = align('center' if center else 'left')
    c.border    = border_thin()
    return c


# ═══════════════════════════════════════════════════════════════════════════════
# GENERAR EXCEL
# ═══════════════════════════════════════════════════════════════════════════════
def generar_excel():
    wb = Workbook()
    wb.remove(wb.active)

    COLOR_PRIORIDAD = {"Alta": ROSA, "Media": NARANJA, "Baja": VERDE}

    # ── Índice General ────────────────────────────────────────────────────────
    ws = wb.create_sheet("Índice General")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 40

    merge_title(ws, "A1:H1",
                "ROSSMIX — Historias de Usuario y Casos de Uso", NEGRO, BLANCO, 15)

    ws.row_dimensions[2].height = 14
    info_items = [
        ("Abono mínimo:", "$5.000 COP"),
        ("Cancelación mínima:", "2 horas antes"),
        ("No asistió:", "Abono no reembolsable"),
        ("Reagenda:", "Abono como crédito"),
    ]
    for i, (lbl, val) in enumerate(info_items, 3):
        ws.cell(row=2, column=i*2-3, value=lbl).font = font(True, 9, GRIS, True)
        ws.cell(row=2, column=i*2-2, value=val).font = font(False, 9, NEGRO)

    ws.row_dimensions[3].height = 8

    headers = ["ID","Título","Actor Principal","Hoja HU","Hoja CU","Prioridad","Pts","Estado"]
    widths  = [8,  42,       20,               12,       12,       12,         6,    14]
    header_row(ws, 4, headers, widths, NEGRO, BLANCO)

    PRIORIDAD_COLOR = {"Alta": ROSA_LIGHT, "Media": NARANJA_L, "Baja": VERDE_L}
    for i, hu in enumerate(HISTORIAS, 5):
        row_bg = GRIS_L if i % 2 == 0 else BLANCO
        p_bg   = PRIORIDAD_COLOR.get(hu[3], BLANCO)
        data_cell(ws, i, 1, hu[0],  row_bg, True, True)
        data_cell(ws, i, 2, hu[1],  row_bg)
        data_cell(ws, i, 3, hu[2],  row_bg)
        data_cell(ws, i, 4, hu[0],  row_bg, False, True)
        cu_id = hu[0].replace("HU","CU")
        data_cell(ws, i, 5, cu_id,  row_bg, False, True)
        pc = ws.cell(row=i, column=6, value=hu[3])
        pc.fill = fill(p_bg); pc.font = font(True,10,ROSA if hu[3]=="Alta" else NARANJA)
        pc.alignment = align('center'); pc.border = border_thin()
        data_cell(ws, i, 7, f"{hu[4]} pts", row_bg, False, True)
        sc = ws.cell(row=i, column=8, value="Implementado")
        sc.fill = fill(VERDE_L); sc.font = font(True,10,VERDE)
        sc.alignment = align('center'); sc.border = border_thin()

    ws.freeze_panes = "A5"

    # ── Sheets HU ────────────────────────────────────────────────────────────
    ACTOR_COLOR = {
        "Cliente": ROSA, "Administrador": NEGRO, "Especialista": MORADO,
        "Cliente/Admin": AZUL, "Cliente/Admin": AZUL,
    }
    for hu in HISTORIAS:
        ws_hu = wb.create_sheet(hu[0])
        ws_hu.sheet_view.showGridLines = False

        ac = next((v for k, v in ACTOR_COLOR.items() if k in hu[2]), GRIS)
        merge_title(ws_hu, "A1:B1", f"{hu[0]} — {hu[1]}", ac, BLANCO, 13)
        ws_hu.row_dimensions[1].height = 35
        ws_hu.column_dimensions["A"].width = 22
        ws_hu.column_dimensions["B"].width = 60

        filas = [
            ("Actor",     hu[2]),
            ("Prioridad", f"{hu[3]}  |  {hu[4]} Story Points"),
            ("Como…",     f"Como {hu[5]}"),
            ("Quiero…",   f"quiero {hu[6]}"),
            ("Para…",     f"para que {hu[7]}"),
            ("Implementación", hu[9]),
        ]
        row = 2
        for lbl, val in filas:
            lc = ws_hu.cell(row=row, column=1, value=lbl)
            lc.fill = fill(GRIS_L); lc.font = font(True,10,GRIS)
            lc.alignment = align(); lc.border = border_thin()
            vc = ws_hu.cell(row=row, column=2, value=val)
            vc.fill = fill(BLANCO); vc.font = font(False,10)
            vc.alignment = align(); vc.border = border_thin()
            ws_hu.row_dimensions[row].height = 28
            row += 1

        ws_hu.cell(row=row, column=1, value="Criterios de Aceptación").fill = fill(ac)
        ws_hu.cell(row=row, column=1).font = font(True,10,BLANCO)
        ws_hu.cell(row=row, column=1).border = border_thin()
        ws_hu.cell(row=row, column=1).alignment = align('center')
        ws_hu.merge_cells(f"B{row}:B{row}")
        row += 1

        for j, crit in enumerate(hu[8], 1):
            lc = ws_hu.cell(row=row, column=1, value=f"CA-{j:02d}")
            lc.fill = fill(ROSA_LIGHT if j%2==0 else BLANCO)
            lc.font = font(True,9,ROSA); lc.alignment = align('center'); lc.border = border_thin()
            vc = ws_hu.cell(row=row, column=2, value=crit)
            vc.fill = fill(ROSA_LIGHT if j%2==0 else BLANCO)
            vc.font = font(False,9); vc.alignment = align(); vc.border = border_thin()
            ws_hu.row_dimensions[row].height = 22
            row += 1

    # ── Sheets CU ────────────────────────────────────────────────────────────
    for cu in CASOS_USO:
        ws_cu = wb.create_sheet(cu[0])
        ws_cu.sheet_view.showGridLines = False
        ws_cu.column_dimensions["A"].width = 24
        ws_cu.column_dimensions["B"].width = 58

        merge_title(ws_cu, "A1:B1", f"{cu[0]} — {cu[1]}", AZUL, BLANCO, 13)
        ws_cu.row_dimensions[1].height = 35

        filas_cu = [
            ("Actor",          cu[2]),
            ("Módulo / Ruta",  cu[3]),
            ("Precondición",   cu[4]),
        ]
        row = 2
        for lbl, val in filas_cu:
            lc = ws_cu.cell(row=row, column=1, value=lbl)
            lc.fill = fill(AZUL_L); lc.font = font(True,10,AZUL)
            lc.alignment = align(); lc.border = border_thin()
            vc = ws_cu.cell(row=row, column=2, value=val)
            vc.fill = fill(BLANCO); vc.font = font(); vc.alignment = align(); vc.border = border_thin()
            ws_cu.row_dimensions[row].height = 24
            row += 1

        ws_cu.cell(row=row, column=1, value="Flujo Principal").fill = fill(AZUL)
        ws_cu.cell(row=row, column=1).font = font(True,10,BLANCO)
        ws_cu.cell(row=row, column=1).border = border_thin()
        ws_cu.cell(row=row, column=1).alignment = align('center')
        row += 1

        for j, paso in enumerate(cu[5], 1):
            lc = ws_cu.cell(row=row, column=1, value=f"Paso {j}")
            lc.fill = fill(AZUL_L if j%2==0 else BLANCO)
            lc.font = font(True,9,AZUL); lc.alignment = align('center'); lc.border = border_thin()
            vc = ws_cu.cell(row=row, column=2, value=paso)
            vc.fill = fill(AZUL_L if j%2==0 else BLANCO)
            vc.font = font(False,9); vc.alignment = align(); vc.border = border_thin()
            ws_cu.row_dimensions[row].height = 22
            row += 1

        ws_cu.cell(row=row, column=1, value="Flujos Alternativos/Excepciones").fill = fill(NARANJA)
        ws_cu.cell(row=row, column=1).font = font(True,10,BLANCO)
        ws_cu.cell(row=row, column=1).border = border_thin()
        ws_cu.cell(row=row, column=1).alignment = align('center')
        row += 1

        for j, alt in enumerate(cu[6], 1):
            lc = ws_cu.cell(row=row, column=1, value=f"Alt-{j}")
            lc.fill = fill(NARANJA_L); lc.font = font(True,9,NARANJA)
            lc.alignment = align('center'); lc.border = border_thin()
            vc = ws_cu.cell(row=row, column=2, value=alt)
            vc.fill = fill(NARANJA_L); vc.font = font(False,9)
            vc.alignment = align(); vc.border = border_thin()
            ws_cu.row_dimensions[row].height = 22
            row += 1

    out_xlsx = os.path.join(OUTDIR, "Rossmix_HU_y_CU_Por_Hojas.xlsx")
    wb.save(out_xlsx)
    print(f"OK: {out_xlsx}")
    return out_xlsx


# ═══════════════════════════════════════════════════════════════════════════════
# GENERAR WORD
# ═══════════════════════════════════════════════════════════════════════════════
def _rgb(hex_str):
    return RGBColor(int(hex_str[0:2],16), int(hex_str[2:4],16), int(hex_str[4:6],16))

def _table_cell(cell, text, bold=False, bg_hex=None, font_color="000000", size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(font_color)
    if bg_hex:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  bg_hex)
        tcPr.append(shd)

def generar_word():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Título principal
    titulo = doc.add_heading("ROSSMIX — Casos de Uso del Sistema", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = _rgb(ROSA)
        run.font.size = Pt(18)

    subtitulo = doc.add_paragraph("Sistema de Gestión de Citas para Salón de Belleza")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitulo.runs:
        run.font.color.rgb = _rgb(GRIS)
        run.font.size = Pt(11)
        run.italic = True
    doc.add_paragraph()

    # Resumen ejecutivo
    doc.add_heading("Resumen Ejecutivo", 1)
    resumen = doc.add_paragraph(
        "Rossmix es un sistema web de gestión de citas para salón de belleza. "
        "Permite a las clientas agendar en línea con un abono de $5.000 COP, "
        "a las especialistas aceptar citas disponibles (modelo tipo Uber) y "
        "a la administradora distribuir, supervisar y gestionar todo el flujo "
        "desde un panel centralizado."
    )
    resumen.style.font.size = Pt(10)
    doc.add_paragraph()

    # Tabla resumen de actores
    doc.add_heading("Actores del Sistema", 2)
    tbl_actores = doc.add_table(rows=1, cols=3)
    tbl_actores.style = 'Table Grid'
    tbl_actores.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Actor", "Tipo de Usuario", "Acceso"]):
        _table_cell(tbl_actores.rows[0].cells[j], h, True, NEGRO, BLANCO, 10, True)

    actores = [
        ("Clienta",         "cliente",      "Portal de citas y perfil personal"),
        ("Administradora",  "admin",        "Panel completo de gestión"),
        ("Especialista",    "especialista", "Portal de aceptación de citas"),
    ]
    for a in actores:
        row = tbl_actores.add_row()
        _table_cell(row.cells[0], a[0], True,  ROSA_LIGHT, ROSA)
        _table_cell(row.cells[1], a[1], False, GRIS_L)
        _table_cell(row.cells[2], a[2], False, BLANCO)
    doc.add_paragraph()

    # Casos de Uso
    doc.add_heading("Especificación de Casos de Uso", 1)

    CU_COLOR = {
        "Cliente": ROSA, "Administrador": NEGRO,
        "Especialista": MORADO, "Cliente/Admin": AZUL,
        "Administrador/Recepcionista": NEGRO,
    }

    for cu in CASOS_USO:
        doc.add_heading(f"{cu[0]} — {cu[1]}", 2)
        for run in doc.paragraphs[-1].runs:
            run.font.color.rgb = _rgb(CU_COLOR.get(cu[2], GRIS))

        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Ancho de columnas
        for row in tbl.rows:
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(12)

        filas_info = [
            ("Actor",         cu[2]),
            ("Módulo / Ruta", cu[3]),
            ("Precondición",  cu[4]),
        ]
        for lbl, val in filas_info:
            r = tbl.add_row()
            _table_cell(r.cells[0], lbl, True, GRIS_L, GRIS)
            _table_cell(r.cells[1], val, False, BLANCO)

        # Flujo principal
        r = tbl.add_row()
        _table_cell(r.cells[0], "Flujo Principal", True, AZUL, BLANCO, center=True)
        tbl.rows[-1].cells[0]._tc.merge(tbl.rows[-1].cells[1]._tc)
        for j, paso in enumerate(cu[5], 1):
            r = tbl.add_row()
            _table_cell(r.cells[0], f"Paso {j}", True,
                        AZUL_L if j%2==0 else BLANCO, AZUL, center=True)
            _table_cell(r.cells[1], paso, False, AZUL_L if j%2==0 else BLANCO)

        # Flujos alternativos
        r = tbl.add_row()
        _table_cell(r.cells[0], "Alternativas / Excepciones", True, NARANJA, BLANCO, center=True)
        tbl.rows[-1].cells[0]._tc.merge(tbl.rows[-1].cells[1]._tc)
        for j, alt in enumerate(cu[6], 1):
            r = tbl.add_row()
            _table_cell(r.cells[0], f"Alt-{j}", True, NARANJA_L, NARANJA, center=True)
            _table_cell(r.cells[1], alt, False, NARANJA_L)

        doc.add_paragraph()

    out_docx = os.path.join(OUTDIR, "Casos de Uso.docx")
    doc.save(out_docx)
    print(f"OK: {out_docx}")
    return out_docx

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n" + "="*60)
    print("  GENERANDO DOCUMENTACIÓN ROSSMIX")
    print("="*60)
    print(f"\nHU totales:  {len(HISTORIAS)}  (originales: 6 + nuevas: {len(HISTORIAS)-6})")
    print(f"CU totales:  {len(CASOS_USO)}  (originales: 6 + nuevos: {len(CASOS_USO)-6})")
    print()
    xlsx = generar_excel()
    docx = generar_word()
    print()
    print(f"Directorio: {OUTDIR}")
    print("="*60)
